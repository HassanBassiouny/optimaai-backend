"""
app/services/bmc_service.py
═══════════════════════════════════════════════════════════════════════════
OPTIMAAI -- BUSINESS MODEL CANVAS GENERATOR (customer-data-aware)
═══════════════════════════════════════════════════════════════════════════

Generates an Osterwalder Business Model Canvas describing the CUSTOMER'S
business, inferred from data they've uploaded into OptimaAi. The output is
ALWAYS a 9-block canvas in classic Osterwalder layout:

    ┌────────────┬────────────┬─────────────┬────────────┬────────────┐
    │ Key        │ Key        │ Value       │ Customer   │ Customer   │
    │ Partners   │ Activities │ Propositions│ Relations  │ Segments   │
    │            ├────────────┤             ├────────────┤            │
    │            │ Key        │             │ Channels   │            │
    │            │ Resources  │             │            │            │
    ├────────────┴────────────┴─────────────┴────────────┴────────────┤
    │     Cost Structure          │      Revenue Streams              │
    └─────────────────────────────┴───────────────────────────────────┘

Public API:
    generate_bmc(business_name: str = "Your Business") -> dict
    get_bmc_file(bmc_id: str, fmt: str) -> Path
"""

from __future__ import annotations

import os
import uuid
from pathlib import Path
from datetime import datetime
from typing import Optional

from app.services.knowledge_base import (
    query_knowledge_base,
    _call_llm,
    kb_stats,
)


# ═══════════════════════════════════════════════════════════════════════════
#   STORAGE
# ═══════════════════════════════════════════════════════════════════════════

BMC_DIR = Path(os.getenv("BMC_DIR", "./optimaai_bmc"))
BMC_DIR.mkdir(parents=True, exist_ok=True)


# ═══════════════════════════════════════════════════════════════════════════
#   THE 9 BMC BLOCKS
# ═══════════════════════════════════════════════════════════════════════════
#
# Each block has a position in the classic Osterwalder grid + a prompt.
# `rag_query` tells us what data to pull from the customer's uploaded
# files for this specific block. Block order matters -- we generate
# Customer Segments first because Value Propositions reference them.
# ───────────────────────────────────────────────────────────────────────────

BMC_BLOCKS: list[dict] = [
    {
        "id":          "customer_segments",
        "title":       "Customer Segments",
        "icon":        "Users",
        "grid_pos":    {"row": 1, "col": 5, "rowspan": 2, "colspan": 1},
        "rag_query":   "customer demographics, regions, repeat customers, order patterns",
        "prompt": (
            "Identify DIFFERENTIATED customer segments. A segment must be "
            "a meaningful subset that's distinct from another subset, with "
            "rough size if the data shows it. Bad segment: 'Customers in "
            "Egypt' (= the whole base). Good segments: 'High-frequency "
            "buyers (3+ orders, e.g. CUST-2205) vs single-purchase "
            "customers', 'Hurghada (X orders) vs Cairo (Y orders)', "
            "'High-AOV customers above the mean of N vs below'. Aim for "
            "2-4 real segments that an operator could target differently."
        ),
    },
    {
        "id":          "value_propositions",
        "title":       "Value Propositions",
        "icon":        "Sparkles",
        "grid_pos":    {"row": 1, "col": 3, "rowspan": 2, "colspan": 1},
        "rag_query":   "products, pricing, offers, top-selling categories, ratings",
        "prompt": (
            "Describe the value the business delivers to its customer "
            "segments. What problem does it solve, what need does it meet, "
            "what is unique about its offering? Anchor each proposition in "
            "actual products, prices, or category breadth from the data."
        ),
    },
    {
        "id":          "channels",
        "title":       "Channels",
        "icon":        "Send",
        "grid_pos":    {"row": 2, "col": 4, "rowspan": 1, "colspan": 1},
        "rag_query":   "payment methods, order channels, fulfillment, delivery",
        "prompt": (
            "Channels = HOW the business reaches and serves customers. NOT "
            "where customers live (that's Customer Segments). Identify "
            "actual channels named in the data: specific payment methods "
            "with their share, fulfillment/delivery types, online/offline "
            "split, marketplace vs direct, etc. If the data doesn't name "
            "specific channels (e.g. only generic 'PaymentMethod' column "
            "with no values), state that the data lacks channel detail "
            "and suggest what to upload."
        ),
    },
    {
        "id":          "customer_relationships",
        "title":       "Customer Relationships",
        "icon":        "Heart",
        "grid_pos":    {"row": 1, "col": 4, "rowspan": 1, "colspan": 1},
        "rag_query":   "customer ratings, reviews, repeat purchases, loyalty, tenure",
        "prompt": (
            "What kind of relationship does the business have with each "
            "customer segment? Self-service, personal assistance, "
            "communities, automated services? Anchor in evidence -- "
            "average tenure, repeat-order rate, rating patterns, etc."
        ),
    },
    {
        "id":          "revenue_streams",
        "title":       "Revenue Streams",
        "icon":        "DollarSign",
        "grid_pos":    {"row": 3, "col": 3, "rowspan": 1, "colspan": 3},
        "rag_query":   "revenue, sales, order value, pricing, transactions, payments",
        "prompt": (
            "List the revenue streams the business actually has, with "
            "concrete figures from the data: average order value, "
            "transaction volume, top revenue contributors, pricing model. "
            "Each stream should be a distinct way money flows in."
        ),
    },
    {
        "id":          "key_resources",
        "title":       "Key Resources",
        "icon":        "Package",
        "grid_pos":    {"row": 2, "col": 2, "rowspan": 1, "colspan": 1},
        "rag_query":   "inventory, products, customer base, technology platform",
        "prompt": (
            "What key resources does the business depend on to deliver its "
            "value proposition? Physical (inventory, fulfillment), "
            "intellectual (brand, customer data), human, financial. "
            "Anchor in observable data: catalog size, customer count, etc."
        ),
    },
    {
        "id":          "key_activities",
        "title":       "Key Activities",
        "icon":        "Activity",
        "grid_pos":    {"row": 1, "col": 2, "rowspan": 1, "colspan": 1},
        "rag_query":   "operations, fulfillment, returns, ratings, transactions",
        "prompt": (
            "What are the most important things the business must do well to "
            "operate? Production, problem-solving, platform maintenance? "
            "Anchor in patterns from the data: return rates, fulfillment "
            "volumes, transaction processing, etc."
        ),
    },
    {
        "id":          "key_partners",
        "title":       "Key Partners",
        "icon":        "Handshake",
        "grid_pos":    {"row": 1, "col": 1, "rowspan": 2, "colspan": 1},
        "rag_query":   "suppliers, vendors, payment processors, logistics, brands",
        "prompt": (
            "Identify ONLY partners that are NAMED in the data: specific "
            "supplier names, specific brand names in the product catalogue, "
            "specific payment processors, specific logistics providers. Do "
            "NOT speculate about generic categories ('banks are key "
            "partners', 'logistics providers are key partners'). If the "
            "data only names brands carried (e.g. Apple) but not vendors "
            "or suppliers, list those brands and add ONE bullet stating "
            "that supplier/vendor records are needed to complete this "
            "block."
        ),
    },
    {
        "id":          "cost_structure",
        "title":       "Cost Structure",
        "icon":        "Receipt",
        "grid_pos":    {"row": 3, "col": 1, "rowspan": 1, "colspan": 2},
        "rag_query":   "costs, expenses, refunds, returns, discounts, margins",
        "prompt": (
            "What are the most important costs in operating this business? "
            "Anchor in real signals visible in the data: refund volumes, "
            "discount levels, return rates, COGS proxies. Identify whether "
            "the business is more cost-driven or value-driven."
        ),
    },
]


# ═══════════════════════════════════════════════════════════════════════════
#   PUBLIC API
# ═══════════════════════════════════════════════════════════════════════════

def generate_bmc(business_name: str = "Your Business") -> dict:
    """
    Generate a full Business Model Canvas inferred from whatever data the
    customer has uploaded into the KB.

    Returns a dict shaped for direct rendering in the frontend grid:

        {
          "status":       "success",
          "bmc_id":       "bmc_abc123",
          "business_name":"Egypt E-commerce",
          "generated_at": "2026-04-25T...",
          "data_scope":   { "sources": [...], "chunks": 7, "is_sparse": False },
          "blocks": [
              { "id":"customer_segments", "title":"...", "icon":"...",
                "grid_pos":{...}, "bullets":[...], "confidence":"medium",
                "data_used":["ecommerce_orders_egypt.csv"] },
              ...
          ],
          "files": { "docx": "...", "pdf": "..." }
        }
    """
    # 1. Probe the KB to know what data we have
    stats   = _safe_kb_stats()
    sources = stats.get("sources", [])
    if not sources:
        return _empty_bmc_response(business_name)

    # 1b. Sample data context to detect currency once. We do this with a
    # broad query so we get a representative sample of all uploaded data.
    sample_ctx, _ = _rag_context_for(
        "currency, region, country, prices, amounts", n_results=2,
    )
    currency = _detect_currency(sample_ctx)
    print(f"  [bmc] Detected currency: {currency}")

    # 1c. Pre-compute behavioral segment dimensions from the data so the
    # Customer Segments block uses them as source of truth.
    segments_hint = _build_segments_hint()
    if segments_hint:
        print(f"  [bmc] Behavioral segments precomputed.")

    # 2. Generate each block
    blocks_out: list[dict] = []
    for blk in BMC_BLOCKS:
        rag_ctx, sources_used = _rag_context_for(blk["rag_query"], n_results=4)
        confidence = _block_confidence(blk["id"], rag_ctx, sources_used, sources)

        bullets = _generate_block_bullets(
            block_title   = blk["title"],
            block_prompt  = blk["prompt"],
            rag_context   = rag_ctx,
            business_name = business_name,
            confidence    = confidence,
            currency      = currency,
            segments_hint = segments_hint if blk["id"] == "customer_segments" else "",
        )

        blocks_out.append({
            "id":         blk["id"],
            "title":      blk["title"],
            "icon":       blk["icon"],
            "grid_pos":   blk["grid_pos"],
            "bullets":    bullets,
            "confidence": confidence,
            "data_used":  sources_used,
        })

    # 3. Persist files
    bmc_id = f"bmc_{uuid.uuid4().hex[:10]}"
    base   = BMC_DIR / bmc_id
    base.mkdir(parents=True, exist_ok=True)

    bmc = {
        "status":        "success",
        "bmc_id":        bmc_id,
        "business_name": business_name,
        "generated_at":  datetime.utcnow().isoformat() + "Z",
        "data_scope": {
            "sources":     sources,
            "chunks":      stats.get("total_chunks", 0),
            "is_sparse":   len(sources) < 2,
        },
        "blocks":        blocks_out,
        "files": {
            "docx": f"/api/v1/bmc/{bmc_id}/file/docx",
            "pdf":  f"/api/v1/bmc/{bmc_id}/file/pdf",
        },
    }

    _render_docx(bmc, base / "bmc.docx")
    _render_pdf (bmc, base / "bmc.pdf")
    return bmc


def get_bmc_file(bmc_id: str, fmt: str) -> Path:
    """Return the on-disk path to a generated BMC artifact."""
    if fmt not in {"docx", "pdf"}:
        raise ValueError(f"Unsupported format: {fmt}")
    path = BMC_DIR / bmc_id / f"bmc.{fmt}"
    if not path.exists():
        raise FileNotFoundError(f"BMC file not found: {bmc_id}.{fmt}")
    return path


# ═══════════════════════════════════════════════════════════════════════════
#   HELPERS
# ═══════════════════════════════════════════════════════════════════════════

def _safe_kb_stats() -> dict:
    """Return KB stats but never throw."""
    try:
        return kb_stats() or {}
    except Exception as e:
        print(f"  [bmc] kb_stats failed: {e}")
        return {}


def _build_segments_hint() -> str:
    """
    Pre-compute behavioral segments from the customer's data and the latest
    ML evaluation. Returns a compact text block that gets injected into
    the Customer Segments prompt as the SOURCE OF TRUTH -- so the LLM
    doesn't invent its own segments.

    Segments built (when data supports them):
      1. Frequency-based (high vs low total_orders)
      2. Engagement-based (avg_days_between_orders thresholds)
      3. Value-based (predicted_clv quartiles or AOV percentiles)
      4. Variety-based (unique_categories)

    If the data is too sparse, returns an empty string -- the prompt then
    falls back to its general segmentation rules.
    """
    pieces: list[str] = []

    # Pull churn-model context which usually has the right features
    churn_ctx, _ = _rag_context_for(
        "churn predictors total_orders avg_days_between_orders "
        "predicted_clv customer_tenure unique_categories",
        n_results=3,
    )
    customer_ctx, _ = _rag_context_for(
        "customer behavior repeat orders frequency CLV high-value",
        n_results=3,
    )
    combined = (churn_ctx + "\n" + customer_ctx).lower()

    # Surface specific features that exist in the data
    feature_hints = []
    if "total_orders" in combined or "total orders" in combined:
        feature_hints.append(
            "- High-frequency buyers (total_orders >= 3) vs single-purchase "
            "customers. Tag this as [derived] when used."
        )
    if ("avg_days_between_orders" in combined
            or "days between" in combined):
        feature_hints.append(
            "- Engaged customers (avg_days_between_orders below median) vs "
            "at-risk customers (avg_days_between_orders above median). Tag "
            "as [derived] when used."
        )
    if "predicted_clv" in combined or "clv" in combined:
        feature_hints.append(
            "- High-CLV customers (top quartile of predicted_clv) vs "
            "average-CLV customers. Tag as [derived] when used."
        )
    if "unique_categories" in combined or "categor" in combined:
        feature_hints.append(
            "- Category-focused buyers (unique_categories <= 1) vs diverse "
            "buyers (unique_categories >= 3). Tag as [derived] when used."
        )

    if not feature_hints:
        return ""

    pieces.append(
        "These behavioral segment dimensions exist in the customer's data "
        "(based on churn-model features and customer-table columns). "
        "Prefer these over geographic-only segments:"
    )
    pieces.extend(feature_hints)
    pieces.append(
        "\nUse the percentage/threshold language above where the data "
        "supports it. If exact percentages are not in the context provided, "
        "describe segments qualitatively and tag them [inferred] instead "
        "of [derived]."
    )
    return "\n".join(pieces)


def _detect_currency(context: str) -> str:
    """
    Determine the currency the customer's data is denominated in. We scan
    for explicit currency tokens AND for region hints. The output is a
    short string like 'EGP', 'USD', or 'the local currency'.

    Scanning explicit tokens first lets us avoid the LLM defaulting to '$'
    when the customer's data clearly says EGP/SAR/AED etc.
    """
    if not context:
        return "the local currency"
    blob = context.lower()
    # Explicit currency codes
    for code in ["egp", "usd", "eur", "gbp", "sar", "aed", "kwd", "qar",
                 "bhd", "omr", "jod", "lbp", "mad", "tnd", "dzd"]:
        if f" {code} " in blob or f" {code}." in blob or f" {code}," in blob \
           or blob.startswith(f"{code} ") or blob.endswith(f" {code}"):
            return code.upper()
    # Region heuristics -- map common region/city hints to likely currency
    region_map = {
        ("egypt", "cairo", "alexandria", "giza", "hurghada"):     "EGP",
        ("uae", "dubai", "abu dhabi", "sharjah"):                 "AED",
        ("saudi", "riyadh", "jeddah", "dammam"):                  "SAR",
        ("kuwait",):                                              "KWD",
        ("qatar", "doha"):                                        "QAR",
        ("bahrain",):                                             "BHD",
        ("oman", "muscat"):                                       "OMR",
        ("usa", "united states", "new york", "los angeles"):      "USD",
        ("uk", "united kingdom", "london"):                       "GBP",
        ("euro", "germany", "france", "spain", "italy"):          "EUR",
    }
    for hints, currency in region_map.items():
        if any(h in blob for h in hints):
            return currency
    return "the local currency"


def _rag_context_for(query: str, n_results: int = 4) -> tuple[str, list[str]]:
    """
    Pull the top-N relevant chunks for a BMC block. Returns (context_text,
    sources_used).
    """
    try:
        result = query_knowledge_base(query, category=None, n_results=n_results)
        if result.get("count", 0) == 0:
            return "", []
        # Filter out latest_ml_results -- BMC is about CUSTOMER data, not
        # platform ML metrics. The ML snapshot is useful for the report
        # generator but irrelevant to a BMC describing the customer's biz.
        sources = [s for s in result.get("sources", []) if s != "latest_ml_results"]
        return result.get("context", ""), sources
    except Exception as e:
        print(f"  [bmc] rag query failed for '{query}': {e}")
        return "", []


def _block_confidence(
    block_id: str,
    rag_ctx: str,
    sources_used: list[str],
    all_sources: list[str],
) -> str:
    """
    Per-block confidence heuristic. Returns 'high' | 'medium' | 'low'.

    Considers:
    - How much context was retrieved
    - How many sources contributed
    - Whether the block is one that's typically NOT covered by transactional
      data alone (Key Partners, Cost Structure, Key Activities tend to be
      under-represented unless the customer uploads supplier/operations data)
    """
    # Block-specific priors. Some blocks are inherently hard to infer from
    # standard ecommerce data even when context retrieves successfully.
    HARD_TO_INFER_FROM_ORDERS = {
        "key_partners",   # need supplier/vendor records
        "key_activities", # need ops data
        "cost_structure", # need P&L / cost data
        "key_resources",  # need asset/team data
    }

    if not rag_ctx or len(rag_ctx) < 200:
        return "low"

    customer_sources = [s for s in all_sources if s != "latest_ml_results"]
    n_used = len(sources_used)
    n_total = len(customer_sources)

    # Base level from retrieval coverage
    if n_used >= 2:
        base = "high"
    elif n_used == 1 and n_total >= 2:
        base = "medium"
    elif n_used == 1:
        base = "medium"
    else:
        base = "low"

    # Demote one level for blocks that need data the customer probably
    # didn't upload. Avoids the "every block is medium" pattern.
    if block_id in HARD_TO_INFER_FROM_ORDERS and n_total <= 1:
        if base == "high":
            base = "medium"
        elif base == "medium":
            base = "low"

    return base


def _generate_block_bullets(
    block_title: str,
    block_prompt: str,
    rag_context: str,
    business_name: str,
    confidence: str,
    currency: str = "the local currency",
    segments_hint: str = "",
) -> list[dict]:
    """
    Call the LLM to produce 2-4 bullets for ONE BMC block. Each bullet now
    returns a dict, not a string:

        { "text": "...", "tag": "observed" | "derived" | "inferred" | "assumption" }

    The tag is the statement-level confidence layer requested by the user.
    The LLM is required to commit to one of these four categories per
    bullet, which forces it to think about whether each claim is grounded.

      observed   = direct fact (a number, a named value, a column from the data)
      derived    = computed from data (a percentage split, a comparison, a top-N)
      inferred   = interpretation that needs judgement; phrased with hedging
                   ("the pattern may indicate...")
      assumption = claim NOT supported by the data; called out explicitly so
                   the reader knows it's not from the upload
    """
    if not rag_context or len(rag_context.strip()) < 100:
        return [{
            "text": f"Insufficient data uploaded to populate {block_title}. "
                    f"Upload data covering this topic to fill this block.",
            "tag": "observed",
        }]

    confidence_clause = {
        "high":   "The data is rich for this block. Most bullets should be "
                  "tagged [observed] or [derived].",
        "medium": "The data partially covers this block. Use [observed] or "
                  "[derived] where you can; mark interpretations clearly "
                  "with [inferred] and any unsupported claims with "
                  "[assumption].",
        "low":    "The data is sparse for this block. Return only the 1-2 "
                  "things the data directly supports as [observed]. Do NOT "
                  "fill space with [inferred] bullets just to reach a count.",
    }[confidence]

    segments_block = ""
    if segments_hint:
        segments_block = (
            "\n\nPRECOMPUTED BEHAVIORAL SEGMENTS FROM THE DATA "
            "(use these directly, do not invent others):\n"
            f"{segments_hint}\n"
        )

    prompt = f"""You are a business strategy consultant filling in ONE block of a
Business Model Canvas for a business called "{business_name}".

The block you are working on: **{block_title}**

YOUR TASK:
{block_prompt}

═══════════════════════════════════════════════════════════════════════════
CRITICAL FORMAT REQUIREMENT (this overrides every other instinct you have):
═══════════════════════════════════════════════════════════════════════════

EVERY bullet must end with one of FOUR evidence tags in square brackets:

  [observed]   = direct fact from data. Use ONLY when you cite a specific
                 number, value, customer ID, product name, or column that
                 appears in the data context below.
                 Example: "Hurghada accounts for 34% of orders [observed]"

  [derived]    = computed / compared from data. Use when you take observed
                 facts and combine them (a ratio, a top-N, a split).
                 Example: "Top-decile customers contribute 24% of revenue [derived]"

  [inferred]   = interpretation requiring judgement. The data points TOWARD
                 this but doesn't prove it. MUST use hedging language:
                 "may indicate", "the pattern is consistent with", "this
                 could reflect". NEVER drop the hedging just because you
                 added the tag.
                 Example: "Repeat purchases above 3 orders may indicate a
                 loyalty segment worth retaining [inferred]"

  [assumption] = claim NOT supported by uploaded data. Use this when you
                 want to mention something a typical business might do
                 but the data does not confirm. Phrase it as an assumption
                 explicitly.
                 Example: "Customer support is assumed to be self-service;
                 the uploaded data contains no support-channel column
                 [assumption]"

═══════════════════════════════════════════════════════════════════════════
HARD RULES:
═══════════════════════════════════════════════════════════════════════════

1. NO UNTAGGED BULLETS. Every line ends with [observed]/[derived]/[inferred]/[assumption].
2. NO GENERIC BUSINESS PHRASES. Forbidden: "managing high-volume order
   fulfillment is associated with efficient operations", "maintaining
   accurate product information is crucial", "providing a seamless customer
   experience". These are vacuous; delete them.
3. NEVER drop hedging on [inferred] bullets. "X is driven by Y [inferred]"
   is wrong because "is driven by" is a confident causal claim. Correct:
   "X may be associated with Y [inferred]".
4. CURRENCY: cite money in {currency} (not $ unless data is USD).
5. CHANNELS = HOW you reach customers (payment methods, fulfillment, online
   vs store). It does NOT mean WHERE customers live. Regions belong in
   Customer Segments.
6. NAMED PARTNERS ONLY. Do not invent partner categories ("banks are key
   partners"). If the data names brands carried (e.g. Apple), list those.
   Otherwise return ONE bullet stating that supplier records are needed,
   tagged [observed].
7. CONFIDENCE GUIDANCE: {confidence_clause}
8. NO META-COMMENTARY. Do not write "here are the bullets" or "based on
   the data" before the list.

═══════════════════════════════════════════════════════════════════════════
OUTPUT FORMAT (strict):
═══════════════════════════════════════════════════════════════════════════

- One bullet per line, each starting with "- ".
- 2 to 4 bullets total. Fewer is fine if the data is sparse.
- Each bullet ends with the evidence tag in square brackets.
- Do NOT use markdown headers, bold, or any other formatting.

═══════════════════════════════════════════════════════════════════════════
DATA CONTEXT FROM CUSTOMER UPLOADS:
═══════════════════════════════════════════════════════════════════════════

{rag_context}{segments_block}

Now write the {block_title} block:"""

    answer, _model, _err = _call_llm(prompt)
    if not answer:
        return [{
            "text": "Generation failed for this block.",
            "tag":  "observed",
        }]

    return _parse_tagged_bullets(answer, block_title)


_VALID_TAGS = {"observed", "derived", "inferred", "assumption"}

# Speculation patterns. Some of these are now ALLOWED inside [inferred]
# bullets because [inferred] explicitly signals interpretation. We still
# block them in [observed] / [derived] bullets where they shouldn't appear.
_SPECULATION_PATTERNS = [
    " likely ", " inferred ", " implied ", " suggests ", " may be ",
    " might be ", " probably ", " possibly ", " appears to ", " seems to ",
    " is presumed ", " could be ", " would be ", " might ", " may ",
]


def _has_speculation(bullet: str) -> bool:
    """Return True if a bullet uses any speculation language. Used by the
    block-level filter (rejects bullets that ignored the tag rule entirely)."""
    padded = " " + bullet.lower() + " "
    return any(p in padded for p in _SPECULATION_PATTERNS)


# Stronger patterns that almost certainly signal an interpretive claim --
# not just a filler word. These are the only patterns that should bump
# [observed]/[derived] bullets to [inferred] in the parser. A bullet with
# concrete numbers should be allowed to use "suggesting X" or "X may be Y"
# as a connecting phrase without losing its observed/derived status.
_STRONG_SPECULATION_PATTERNS = [
    " is likely ", " are likely ", " was likely ", " were likely ",
    " is inferred ", " are inferred ", " is presumed ", " are presumed ",
    " is implied ", " are implied ",
    " probably is ", " probably are ",
    " appears to be ", " appears to ", " seems to be ", " seems to ",
    "likely indicates", "likely reflects", "is presumably",
]


def _has_strong_speculation(bullet: str) -> bool:
    """Return True only for explicit interpretive phrasing -- not for
    filler words like 'may' or 'suggests' which can legitimately appear
    in observed/derived claims."""
    padded = " " + bullet.lower() + " "
    return any(p in padded for p in _STRONG_SPECULATION_PATTERNS)


_TAG_REGEX = None  # lazily compiled


def _strip_all_tags(text: str) -> tuple[str, Optional[str]]:
    """
    Find every `[tag]` marker in the text, return the cleaned text
    (with all tags removed) and the LAST recognised tag found. The
    LLM sometimes emits multiple tags ('the value is X [observed],
    suggesting Y [inferred]') -- we trust the last one as the
    bullet's overall classification.
    """
    import re
    global _TAG_REGEX
    if _TAG_REGEX is None:
        _TAG_REGEX = re.compile(r"\[([a-zA-Z]+)\]")

    found_tag: Optional[str] = None
    cleaned = text
    for m in _TAG_REGEX.finditer(text):
        candidate = m.group(1).lower()
        if candidate in _VALID_TAGS:
            found_tag = candidate  # last match wins
    # Strip every recognised-tag bracket (leave non-tag brackets intact
    # in case the LLM uses `[note]` for something legitimate)
    cleaned = _TAG_REGEX.sub(
        lambda m: "" if m.group(1).lower() in _VALID_TAGS else m.group(0),
        text,
    )
    # Tidy double spaces / leading punctuation left by the strip
    cleaned = " ".join(cleaned.split())
    # Strip trailing ellipsis or other unfinished-sentence markers, then
    # add a clean period
    cleaned = cleaned.rstrip(" .,;:…").rstrip()
    if cleaned and not cleaned.endswith((".", "!", "?")):
        cleaned += "."
    return cleaned, found_tag


def _parse_tagged_bullets(raw: str, block_title: str) -> list[dict]:
    """
    Parse the LLM's tagged-bullet output into a list of {text, tag} dicts.

    The parser tolerates several forms of LLM misbehaviour:
      - tag in the middle of the bullet ("X = N [observed], suggesting Y")
      - multiple tags in one bullet (uses the last)
      - missing tag entirely (defaults to [inferred])
      - capital-cased or stray-whitespace tags
    """
    bullets: list[dict] = []
    for line in raw.strip().split("\n"):
        line = line.strip()
        if not line:
            continue
        # Strip leading bullet markers / numbers
        for prefix in ("- ", "* ", "• "):
            if line.startswith(prefix):
                line = line[len(prefix):].strip()
                break
        else:
            if len(line) > 2 and line[0].isdigit() and line[1] in ".):":
                line = line[2:].strip()

        if not line or len(line) < 5:
            continue

        # Strip any/all [tag] markers, keeping the last recognised one
        text, tag = _strip_all_tags(line)

        if tag is None:
            # No tag found anywhere -- default to [inferred]
            tag = "inferred"

        # Auto-bump observed/derived bullets that contain heavy speculation
        # language (the tag must match the language). Note: this is now
        # narrower -- we only bump for STRONG speculation, not for filler
        # words like "may" or "suggests" that often appear in observed
        # claims as connecting language.
        if tag in ("observed", "derived") and _has_strong_speculation(text):
            tag = "inferred"

        bullets.append({"text": text, "tag": tag})

    if not bullets:
        return [{
            "text": (f"The uploaded data does not directly support claims "
                     f"about {block_title}. Upload more relevant data to "
                     "populate this block."),
            "tag":  "observed",
        }]

    return bullets[:4]


def _empty_bmc_response(business_name: str) -> dict:
    """When no customer data has been uploaded yet, return an empty shell."""
    return {
        "status":        "no_data",
        "business_name": business_name,
        "generated_at":  datetime.utcnow().isoformat() + "Z",
        "blocks":        [],
        "data_scope":    {"sources": [], "chunks": 0, "is_sparse": True},
        "message":       (
            "No customer data found in the knowledge base. Upload at least "
            "one CSV through the Datasets page to generate a Business Model "
            "Canvas that describes your business."
        ),
    }


# ═══════════════════════════════════════════════════════════════════════════
#   RENDERING — DOCX + PDF
# ═══════════════════════════════════════════════════════════════════════════

def _render_docx(bmc: dict, out_path: Path) -> None:
    """Render the BMC to a Word document with a 9-block table."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_ALIGN_VERTICAL
    except ImportError:
        print("  [bmc] python-docx not installed; falling back to plain text")
        out_path.write_text(_bmc_to_markdown(bmc), encoding="utf-8")
        return

    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Set landscape so the canvas grid fits properly
    section = doc.sections[0]
    new_w, new_h = section.page_height, section.page_width
    section.orientation = 1  # landscape
    section.page_width  = new_w
    section.page_height = new_h
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(f"Business Model Canvas — {bmc['business_name']}")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        f"OptimaAi  ·  Generated {datetime.utcnow().strftime('%B %d, %Y')}"
    )
    sub_run.italic = True
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # The Osterwalder layout in tabular form. We render it as a sequence of
    # bordered tables since reproducing the exact 5-col-3-row spanning grid
    # in python-docx is fiddly. The order below matches the standard
    # left-to-right, top-to-bottom reading order.
    blocks_by_id = {b["id"]: b for b in bmc["blocks"]}
    layout_rows = [
        ["key_partners", "key_activities", "value_propositions",
         "customer_relationships", "customer_segments"],
        ["key_partners", "key_resources", "value_propositions",
         "channels", "customer_segments"],  # spans handled below
        ["cost_structure", "cost_structure", "revenue_streams",
         "revenue_streams", "revenue_streams"],
    ]

    table = doc.add_table(rows=3, cols=5)
    table.style = "Table Grid"
    table.autofit = False

    seen_in_row = [set(), set(), set()]
    for r, row in enumerate(layout_rows):
        for c, blk_id in enumerate(row):
            cell = table.cell(r, c)
            blk  = blocks_by_id.get(blk_id)
            if not blk:
                continue
            # Avoid re-writing into a cell already merged from above
            already = (r, c) in {
                (rr, cc) for rr in range(r) for cc in range(5)
                if layout_rows[rr][cc] == blk_id and rr != r
            }
            if blk_id in seen_in_row[r]:
                continue  # cell already merged horizontally
            seen_in_row[r].add(blk_id)

            # Merge horizontally with neighbours that share blk_id in same row
            cspan = 1
            while c + cspan < 5 and layout_rows[r][c + cspan] == blk_id:
                table.cell(r, c).merge(table.cell(r, c + cspan))
                cspan += 1
            # Merge vertically with cells below that share blk_id in same col
            for rr in range(r + 1, 3):
                if layout_rows[rr][c] == blk_id:
                    table.cell(r, c).merge(table.cell(rr, c))

            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.text = ""  # clear default
            # Header
            hdr = cell.paragraphs[0]
            hdr_run = hdr.add_run(blk["title"])
            hdr_run.bold = True
            hdr_run.font.size = Pt(11)
            hdr_run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)
            # Confidence badge
            conf_run = hdr.add_run(f"  ({blk['confidence']})")
            conf_run.italic = True
            conf_run.font.size = Pt(8)
            conf_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            # Bullets — each is now a {text, tag} dict
            for b in blk["bullets"]:
                if isinstance(b, dict):
                    text, tag = b.get("text", ""), b.get("tag", "observed")
                else:
                    text, tag = str(b), "observed"
                p = cell.add_paragraph(style="List Bullet")
                br = p.add_run(text + " ")
                br.font.size = Pt(9)
                # Tag suffix in a tinted color matching the tag type
                tag_run = p.add_run(f"[{tag}]")
                tag_run.italic = True
                tag_run.font.size = Pt(7)
                hex_color = _TAG_HEX.get(tag, "6B7280")
                tag_run.font.color.rgb = RGBColor(
                    int(hex_color[0:2], 16),
                    int(hex_color[2:4], 16),
                    int(hex_color[4:6], 16),
                )

    doc.save(out_path)


def _render_pdf(bmc: dict, out_path: Path) -> None:
    """Render the BMC to a single-page A3 landscape PDF.

    Sizing rationale: A3 landscape = 16.54 x 11.69 in. With 0.4 in margins
    we have 15.74 x 10.89 in usable. Title block ~0.7 in -> grid gets
    ~10.1 in vertical space in 3 rows = 3.4 in each. That comfortably fits
    a header line + confidence line + 3 bullets per block.

    A3 is what Strategyzer's printable BMC templates use; Letter landscape
    just doesn't have enough vertical room for a 9-block canvas with prose.
    The Word doc and the on-screen UI keep the full content; this PDF is
    the at-a-glance executive view.
    """
    try:
        from reportlab.lib.pagesizes import landscape, A3
        from reportlab.lib.units    import inch
        from reportlab.lib          import colors
        from reportlab.platypus     import (
            SimpleDocTemplate, Paragraph, Table, TableStyle, KeepTogether,
        )
        from reportlab.lib.styles   import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums    import TA_CENTER
    except ImportError:
        print("  [bmc] reportlab not installed; falling back to plain text")
        out_path.write_text(_bmc_to_markdown(bmc), encoding="utf-8")
        return

    # ── Styles ─────────────────────────────────────────────────────────
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "BMCTitle", parent=styles["Title"],
        fontSize=20, leading=24,
        textColor=colors.HexColor("#1F2A44"),
        alignment=TA_CENTER, spaceAfter=2,
    )
    sub_style = ParagraphStyle(
        "BMCSub", parent=styles["Italic"],
        fontSize=10, leading=12,
        textColor=colors.HexColor("#6B7280"),
        alignment=TA_CENTER, spaceAfter=14,
    )
    block_title_style = ParagraphStyle(
        "BlockTitle", parent=styles["Heading4"],
        fontSize=11, leading=13, fontName="Helvetica-Bold",
        textColor=colors.HexColor("#1F2A44"), spaceAfter=2,
    )
    bullet_style = ParagraphStyle(
        "Bullet", parent=styles["BodyText"],
        fontSize=8.5, leading=11,
        textColor=colors.HexColor("#111827"),
        leftIndent=8, bulletIndent=0, spaceAfter=2,
    )
    conf_style = ParagraphStyle(
        "Conf", parent=styles["Italic"],
        fontSize=7.5, leading=9,
        textColor=colors.HexColor("#6B7280"), spaceAfter=4,
    )

    # ── Document (A3 landscape) ────────────────────────────────────────
    page_w, page_h = landscape(A3)
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=landscape(A3),
        leftMargin=0.4 * inch, rightMargin=0.4 * inch,
        topMargin=0.4 * inch,  bottomMargin=0.4 * inch,
        title=f"BMC — {bmc['business_name']}", author="OptimaAi",
    )

    story: list = []

    # ── Helpers ────────────────────────────────────────────────────────
    blocks_by_id = {b["id"]: b for b in bmc["blocks"]}

    # Per-block bullet caps. Top-row blocks have narrow columns so they
    # need fewer bullets; bottom-row blocks span wider columns and have
    # less vertical space available.
    BULLET_CAPS = {
        "key_partners":           3,
        "key_activities":         2,
        "value_propositions":     3,
        "customer_relationships": 2,
        "customer_segments":      3,
        "key_resources":          2,
        "channels":               2,
        "cost_structure":         3,   # wide column, can fit 3
        "revenue_streams":        3,   # wide column, can fit 3
    }

    def _truncate(text: str, max_chars: int = 200) -> str:
        """Cap each bullet so cells stay scannable. A3 cells comfortably
        fit ~200 chars per bullet at 8.5pt. End with a period when we have
        to cut, not an ellipsis -- ellipses make it look like the tag-strip
        process broke something."""
        text = (text or "").strip()
        if len(text) <= max_chars:
            return text
        cut = text.rfind(" ", 0, max_chars)
        if cut < max_chars - 40:
            cut = max_chars
        return text[:cut].rstrip(",.;:") + "."

    def render_block(blk_id: str):
        blk = blocks_by_id.get(blk_id)
        if not blk:
            return [Paragraph("&nbsp;", bullet_style)]
        flow = [
            Paragraph(blk["title"], block_title_style),
            Paragraph(f"({blk['confidence']} confidence)", conf_style),
        ]
        cap = BULLET_CAPS.get(blk_id, 2)
        for b in blk["bullets"][:cap]:
            if isinstance(b, dict):
                text, tag = b.get("text", ""), b.get("tag", "observed")
            else:
                text, tag = str(b), "observed"
            color = _TAG_HEX.get(tag, "6B7280")
            # Inline tag with HTML font tag (reportlab Paragraph supports it)
            flow.append(Paragraph(
                f"• {_truncate(text)} "
                f'<font size="6.5" color="#{color}"><i>[{tag}]</i></font>',
                bullet_style,
            ))
        return flow

    # ── Header row content (title + subtitle) ─────────────────────────
    # Embedding the header inside the grid table guarantees it stays with
    # the canvas on the same page. Otherwise KeepTogether tends to push
    # the entire grid to page 2 if the title block consumes too much
    # height up top.
    header_cell = [
        Paragraph(
            f"Business Model Canvas — {bmc['business_name']}", title_style,
        ),
        Paragraph(
            f"OptimaAi  &middot;  Generated "
            f"{datetime.utcnow().strftime('%B %d, %Y')}",
            sub_style,
        ),
    ]

    # ── Grid geometry ─────────────────────────────────────────────────
    # Usable area: 16.54 x 11.69 - 0.8 margins = 15.74 x 10.89 in.
    # Header row gets 0.6 in. Remaining 10.29 in split between 3 grid rows
    # with a 0.2 in safety margin so floating-point rounding doesn't push
    # the table 0.01 in over the page boundary and trigger a split.
    usable_w   = page_w - 0.8 * inch
    col_w      = usable_w / 5
    grid_total = page_h - 0.8 * inch
    header_h   = 0.6 * inch
    safety     = 0.3 * inch                     # buffer against rounding
    grid_h     = grid_total - header_h - safety
    top_row_h    = grid_h * 0.40                # ~3.85 in each
    bottom_row_h = grid_h * 0.20                # ~1.93 in

    # ── 5x4 grid: 1 header row + 3 canvas rows ────────────────────────
    grid_data = [
        # header row -- spans all 5 columns
        [header_cell, "", "", "", ""],
        # row 1
        [render_block("key_partners"), render_block("key_activities"),
         render_block("value_propositions"), render_block("customer_relationships"),
         render_block("customer_segments")],
        # row 2 -- key_partners / value_props / customer_segments span down
        ["", render_block("key_resources"),
         "", render_block("channels"), ""],
        # row 3 -- cost spans cols 0-1, revenue spans cols 2-4
        [render_block("cost_structure"), "",
         render_block("revenue_streams"), "", ""],
    ]

    t = Table(
        grid_data,
        colWidths=[col_w] * 5,
        rowHeights=[header_h, top_row_h, top_row_h, bottom_row_h],
    )
    t.setStyle(TableStyle([
        ("VALIGN",        (0, 1), (-1, -1), "TOP"),
        ("VALIGN",        (0, 0), (-1, 0),  "MIDDLE"),
        ("ALIGN",         (0, 0), (-1, 0),  "CENTER"),
        ("BOX",           (0, 1), (-1, -1), 1.2, colors.HexColor("#1F2A44")),
        ("INNERGRID",     (0, 1), (-1, -1), 0.5, colors.HexColor("#D1D5DB")),
        ("LEFTPADDING",   (0, 1), (-1, -1), 8),
        ("RIGHTPADDING",  (0, 1), (-1, -1), 8),
        ("TOPPADDING",    (0, 1), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 1), (-1, -1), 8),
        # Header spans all 5 columns
        ("SPAN", (0, 0), (4, 0)),
        # Vertical spans across rows 1-2 (the tall side blocks)
        ("SPAN", (0, 1), (0, 2)),  # Key Partners
        ("SPAN", (2, 1), (2, 2)),  # Value Propositions
        ("SPAN", (4, 1), (4, 2)),  # Customer Segments
        # Horizontal spans in row 3 (the bottom row)
        ("SPAN", (0, 3), (1, 3)),  # Cost Structure
        ("SPAN", (2, 3), (4, 3)),  # Revenue Streams
        # Subtle accent shading on the bottom row
        ("BACKGROUND", (0, 3), (1, 3), colors.HexColor("#FEF3F2")),  # cost
        ("BACKGROUND", (2, 3), (4, 3), colors.HexColor("#ECFDF5")),  # revenue
    ]))

    story.append(t)
    doc.build(story)


def _bmc_to_markdown(bmc: dict) -> str:
    """Plain-text fallback for the doc/pdf renderers."""
    lines = [
        f"# Business Model Canvas — {bmc['business_name']}",
        f"_Generated by OptimaAi on "
        f"{datetime.utcnow().strftime('%B %d, %Y')}_",
        "",
    ]
    for blk in bmc["blocks"]:
        lines += [
            f"## {blk['title']}  _( {blk['confidence']} confidence )_",
            "",
        ]
        for b in blk["bullets"]:
            if isinstance(b, dict):
                text, tag = b.get("text", ""), b.get("tag", "observed")
                lines.append(f"- {text} _[{tag}]_")
            else:
                lines.append(f"- {b}")
        lines.append("")
    return "\n".join(lines)


# Tag colors used by both DOCX and PDF renderers. Maps the four evidence
# tags to hex colors so a reader can scan a block and see at a glance
# how grounded each bullet is.
_TAG_HEX = {
    "observed":   "047857",   # emerald-700  — direct fact
    "derived":    "1D4ED8",   # blue-700     — computed
    "inferred":   "B45309",   # amber-700    — interpretation
    "assumption": "BE123C",   # rose-700     — not in data
}