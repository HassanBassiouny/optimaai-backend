"""
app/services/report_service.py
═══════════════════════════════════════════════════════════════════════════
OPTIMAAI -- AI REPORT GENERATOR
═══════════════════════════════════════════════════════════════════════════

Generates structured, multi-section business reports by combining:
  - Trained ML model outputs (revenue, churn, growth)
  - Uploaded data context retrieved via RAG
  - LLM-written narrative for each section

Templates:
  - executive_summary  : top-level overview for leadership
  - churn_deep_dive    : focused churn analysis with drivers and actions
  - revenue_forecast   : revenue model performance + 12-period projection
  - growth_outlook     : growth model + monthly projections

Outputs:
  - markdown (str)            for inline UI rendering
  - .docx file                downloadable Word document
  - .pdf  file                downloadable PDF (rendered from the same docx)

Public API:
    generate_report(template_id: str, role: str = "executive") -> dict
    list_templates() -> list[dict]
    get_report_file(report_id: str, fmt: str) -> Path
"""

from __future__ import annotations

import os
import io
import json
import uuid
import tempfile
from pathlib import Path
from datetime import datetime
from typing import Optional

# ── Local services ─────────────────────────────────────────────────────────
from app.services.ml_bridge import load_latest_evaluation
from app.services.knowledge_base import (
    query_knowledge_base,
    _call_llm,                        # private but stable -- already in use elsewhere
    _format_revenue_ml,
    _format_churn_ml,
    _format_growth_ml,
)

# ═══════════════════════════════════════════════════════════════════════════
#   STORAGE
# ═══════════════════════════════════════════════════════════════════════════

REPORTS_DIR = Path(os.getenv("REPORTS_DIR", "./optimaai_reports"))
REPORTS_DIR.mkdir(parents=True, exist_ok=True)


# ═══════════════════════════════════════════════════════════════════════════
#   TEMPLATES
# ═══════════════════════════════════════════════════════════════════════════
#
# Each template is a list of sections. Each section declares:
#   - id          unique key
#   - title       printed heading
#   - prompt      LLM instruction to write the section
#   - ml_data     what ML context to feed the prompt
#                 ("revenue" | "churn" | "growth" | "all" | None)
#   - rag_query   optional RAG question to retrieve uploaded-data context
# ───────────────────────────────────────────────────────────────────────────

_TEMPLATES: dict[str, dict] = {
    "executive_summary": {
        "title":       "Executive Summary",
        "description": "Top-level overview for leadership: revenue, churn, "
                       "growth, model accuracy, and three priority actions.",
        "icon":        "FileText",
        "sections": [
            {
                "id":        "headline",
                "title":     "Headline",
                "prompt":    ("Write a single-paragraph executive headline (4-6 sentences) "
                              "summarising the overall business position based on the "
                              "ML model results below. Lead with the most important number."),
                "ml_data":   "all",
                "rag_query": None,
            },
            {
                "id":        "revenue",
                "title":     "Revenue Performance",
                "prompt":    ("Write a 2-paragraph analysis of the revenue forecasting model. "
                              "Paragraph 1: model accuracy and what it predicts for the next "
                              "12 periods. Paragraph 2: the biggest revenue drivers and what "
                              "they mean for the business."),
                "ml_data":   "revenue",
                "rag_query": "revenue trends and sales drivers",
            },
            {
                "id":        "churn",
                "title":     "Customer Churn Risk",
                "prompt":    ("Write a 2-paragraph analysis of customer churn. Paragraph 1: "
                              "the current churn rate and how reliably the model can identify "
                              "at-risk customers. Paragraph 2: the top behavioural predictors "
                              "of churn and what they signal."),
                "ml_data":   "churn",
                "rag_query": "customer churn and retention",
            },
            {
                "id":        "growth",
                "title":     "Growth Outlook",
                "prompt":    ("Write a 2-paragraph analysis of growth. Paragraph 1: the next "
                              "12-month revenue trajectory with specific monthly figures. "
                              "Paragraph 2: what's driving growth and any seasonality signals "
                              "in the data."),
                "ml_data":   "growth",
                "rag_query": "growth and monthly trends",
            },
            {
                "id":        "actions",
                "title":     "Priority Actions",
                "prompt":    ("Based on the ML results below, list exactly THREE concrete, "
                              "numbered priority actions for leadership in the next 90 days. "
                              "Each action: 2 sentences. Anchor each one to a specific number "
                              "from the model results."),
                "ml_data":   "all",
                "rag_query": None,
            },
        ],
    },

    "churn_deep_dive": {
        "title":       "Customer Churn Deep-Dive",
        "description": "Focused churn analysis: drivers, risk segments, "
                       "retention recommendations.",
        "icon":        "Users",
        "sections": [
            {
                "id":        "overview",
                "title":     "Churn Landscape",
                "prompt":    ("Write a 2-paragraph overview of the current churn situation. "
                              "Use the model's churn rate, sample size, and ROC-AUC to "
                              "characterise both the scale of the problem and the "
                              "reliability of the model's risk estimates."),
                "ml_data":   "churn",
                "rag_query": "customer churn and retention",
            },
            {
                "id":        "drivers",
                "title":     "Top Churn Drivers",
                "prompt":    ("Walk through the TOP FIVE churn predictors one by one. "
                              "For each, give the feature name, its importance score, and "
                              "a 1-2 sentence interpretation of what it tells you about "
                              "customer behaviour."),
                "ml_data":   "churn",
                "rag_query": None,
            },
            {
                "id":        "model_performance",
                "title":     "Model Performance",
                "prompt":    ("Write one paragraph about how reliable the churn model is. "
                              "Reference ROC-AUC, F1, precision, and recall. Translate the "
                              "metrics into plain English for a non-technical reader."),
                "ml_data":   "churn",
                "rag_query": None,
            },
            {
                "id":        "recommendations",
                "title":     "Retention Recommendations",
                "prompt":    ("Provide FIVE concrete, numbered retention actions. Each action "
                              "must reference a specific top-driver feature from the model. "
                              "Format: '1. Action title -- 2-sentence description'."),
                "ml_data":   "churn",
                "rag_query": "customer demographics and segments",
            },
        ],
    },

    "revenue_forecast": {
        "title":       "Revenue Forecast Report",
        "description": "Revenue model accuracy, drivers, and 12-period projection.",
        "icon":        "TrendingUp",
        "sections": [
            {
                "id":        "summary",
                "title":     "Forecast Summary",
                "prompt":    ("Write a 2-paragraph summary. Paragraph 1: state the next-12-period "
                              "average and range with specific numbers. Paragraph 2: comment "
                              "on whether the trajectory is flat, rising, or falling, and "
                              "what that implies."),
                "ml_data":   "revenue",
                "rag_query": "revenue and sales trends",
            },
            {
                "id":        "accuracy",
                "title":     "Model Accuracy",
                "prompt":    ("Write one paragraph about the model's accuracy. Reference MAPE, "
                              "R-squared, RMSE, and forecast bias. Translate the technical "
                              "metrics into plain English."),
                "ml_data":   "revenue",
                "rag_query": None,
            },
            {
                "id":        "drivers",
                "title":     "Revenue Drivers",
                "prompt":    ("Walk through the TOP FIVE revenue drivers one by one. For each, "
                              "give the feature name, its importance score, and a 1-2 sentence "
                              "interpretation of why it matters."),
                "ml_data":   "revenue",
                "rag_query": "product and pricing data",
            },
            {
                "id":        "actions",
                "title":     "Strategic Recommendations",
                "prompt":    ("Based on the top revenue drivers and the forecast trajectory, "
                              "list THREE numbered strategic recommendations. Each: 2 sentences, "
                              "anchored to specific numbers."),
                "ml_data":   "revenue",
                "rag_query": None,
            },
        ],
    },

    "growth_outlook": {
        "title":       "Growth Outlook Report",
        "description": "12-month growth forecast with seasonality, drivers, "
                       "and strategic implications.",
        "icon":        "BarChart3",
        "sections": [
            {
                "id":        "summary",
                "title":     "Growth Trajectory",
                "prompt":    ("Write a 2-paragraph summary of the 12-month growth forecast. "
                              "Paragraph 1: cite specific monthly figures and identify the "
                              "peak month and trough month. Paragraph 2: characterise the "
                              "overall direction and any obvious seasonality."),
                "ml_data":   "growth",
                "rag_query": "growth metrics and seasonality",
            },
            {
                "id":        "drivers",
                "title":     "Growth Drivers",
                "prompt":    ("Walk through the TOP FIVE growth drivers from the model. For each, "
                              "give the feature, its importance score, and 1-2 sentences of "
                              "interpretation."),
                "ml_data":   "growth",
                "rag_query": None,
            },
            {
                "id":        "model_performance",
                "title":     "Model Reliability",
                "prompt":    ("Write one paragraph about the growth model's reliability. "
                              "Reference MAPE, R-squared, RMSE, and the train/test split. "
                              "Translate the metrics into plain English."),
                "ml_data":   "growth",
                "rag_query": None,
            },
            {
                "id":        "actions",
                "title":     "Strategic Implications",
                "prompt":    ("Based on the trajectory and drivers, list THREE numbered "
                              "strategic recommendations for capitalising on the forecast. "
                              "Each: 2 sentences, anchored to a specific monthly figure or "
                              "driver."),
                "ml_data":   "growth",
                "rag_query": None,
            },
        ],
    },
}


# ═══════════════════════════════════════════════════════════════════════════
#   PUBLIC API -- listing
# ═══════════════════════════════════════════════════════════════════════════

def list_templates() -> list[dict]:
    """List all available report templates for the frontend to render."""
    return [
        {
            "id":           tid,
            "title":        tpl["title"],
            "description":  tpl["description"],
            "icon":         tpl["icon"],
            "section_count": len(tpl["sections"]),
        }
        for tid, tpl in _TEMPLATES.items()
    ]


# ═══════════════════════════════════════════════════════════════════════════
#   PUBLIC API -- generation
# ═══════════════════════════════════════════════════════════════════════════

def generate_report(template_id: str, role: str = "executive") -> dict:
    """
    Generate a report from one of the registered templates.

    Returns a dict:
        {
          "report_id":  "rep_abc123",
          "template":   "executive_summary",
          "title":      "...",
          "generated_at": iso8601,
          "sections":   [{ "id", "title", "content_md" }, ...],
          "files":      { "docx": "/api/v1/reports/<id>/docx",
                          "pdf":  "/api/v1/reports/<id>/pdf" },
          "markdown":   "...full report as markdown..."
        }
    """
    if template_id not in _TEMPLATES:
        raise ValueError(f"Unknown template: {template_id}")
    tpl = _TEMPLATES[template_id]

    # Pull ML KPIs once -- avoids re-reading the JSON for every section
    try:
        kpis = load_latest_evaluation()
    except Exception as e:
        print(f"  [report] No ML artefacts available: {e}")
        kpis = {}

    # ── Data-quality sanity check ──────────────────────────────────────
    # Programmatic check that runs BEFORE the LLM. Catches statistically
    # implausible metrics so they are surfaced even if the LLM doesn't.
    warnings = _data_quality_warnings(kpis)

    # ── Confidence scoring ─────────────────────────────────────────────
    # Programmatic confidence rating for each model (high / medium / low)
    # based on training size, metric quality, and detected warnings.
    # Surfaced both in the warnings banner AND fed into every LLM prompt.
    confidence    = _model_confidence(kpis, warnings)
    confidence_md = _format_confidence_md(confidence)
    gates_md      = _format_decision_gates_md(confidence)

    sections_out = []

    # ── Section 1: DECISION GATE SCORECARD ────────────────────────────
    # At-a-glance verdict for every business decision. Lands at the top
    # of the report so leadership knows what to act on without reading
    # the body. Built programmatically -- the LLM cannot override it.
    if gates_md:
        sections_out.append({
            "id":         "decision_gates",
            "title":      "🎯 Decision Gates",
            "content_md": gates_md,
        })

    # ── Section 2: DATA QUALITY & CONFIDENCE ──────────────────────────
    # The detail behind the gates: WHY each model got the rating it did,
    # plus any data-quality concerns the LLM should not gloss over.
    opener_parts = []
    if warnings:
        opener_parts.append(
            "_The following data-quality concerns were detected in your model "
            "evaluation results. Treat the rest of this report with appropriate "
            "caution until they are resolved._\n\n"
            + "\n".join(f"- {w}" for w in warnings)
        )
    if confidence_md:
        opener_parts.append(confidence_md)

    if opener_parts:
        sections_out.append({
            "id":         "data_quality_warnings",
            "title":      "⚠ Data Quality & Confidence",
            "content_md": "\n\n".join(opener_parts),
        })
    # Two-pass generation:
    #   Pass 1 -- write every section EXCEPT synthesis sections (actions /
    #             recommendations / priority) using just ML + RAG context.
    #   Pass 2 -- write synthesis sections AFTER, with the entire previous
    #             draft folded into the prompt so recommendations can be
    #             grounded in the actual analysis rather than raw metrics.
    SYNTHESIS_IDS = {"actions", "recommendations", "priority_actions"}

    pass1_sections: list[dict] = []
    pass2_specs:    list[dict] = []

    for sec in tpl["sections"]:
        if sec["id"] in SYNTHESIS_IDS:
            pass2_specs.append(sec)
            continue

        ml_ctx  = _ml_context_for(sec["ml_data"], kpis)
        rag_ctx = _rag_context_for(sec["rag_query"]) if sec["rag_query"] else ""

        prompt = _build_section_prompt(
            section_title = sec["title"],
            instructions  = sec["prompt"],
            role          = role,
            ml_context    = ml_ctx,
            rag_context   = rag_ctx,
            confidence    = confidence,
        )
        answer, _model, _err = _call_llm(prompt)
        content = (answer or "").strip() or "_(generation failed for this section)_"
        content = _strip_duplicate_heading(content, sec["title"])
        pass1_sections.append({
            "id":         sec["id"],
            "title":      sec["title"],
            "content_md": content,
        })

    # Build a digest of pass-1 content so synthesis sections can read it
    pass1_digest = "\n\n".join(
        f"### {s['title']}\n{s['content_md']}" for s in pass1_sections
    )

    pass2_sections: list[dict] = []
    for sec in pass2_specs:
        ml_ctx  = _ml_context_for(sec["ml_data"], kpis)
        rag_ctx = _rag_context_for(sec["rag_query"]) if sec["rag_query"] else ""

        prompt = _build_synthesis_prompt(
            section_title = sec["title"],
            instructions  = sec["prompt"],
            role          = role,
            ml_context    = ml_ctx,
            rag_context   = rag_ctx,
            prior_draft   = pass1_digest,
            warnings      = warnings,
            confidence    = confidence,
        )
        answer, _model, _err = _call_llm(prompt)
        content = (answer or "").strip() or "_(generation failed for this section)_"
        content = _strip_duplicate_heading(content, sec["title"])
        pass2_sections.append({
            "id":         sec["id"],
            "title":      sec["title"],
            "content_md": content,
        })

    # Combine in original template order: warnings first, then merge
    # pass1 + pass2 back together preserving the template's intended order.
    by_id = {s["id"]: s for s in (pass1_sections + pass2_sections)}
    sections_out.extend(by_id[sec["id"]] for sec in tpl["sections"] if sec["id"] in by_id)

    # Stitch the sections into a single markdown document
    full_md = _stitch_markdown(tpl["title"], sections_out)

    # Persist three artefacts: .md, .docx, .pdf
    report_id = f"rep_{uuid.uuid4().hex[:10]}"
    base      = REPORTS_DIR / report_id
    base.mkdir(parents=True, exist_ok=True)
    (base / "report.md").write_text(full_md, encoding="utf-8")

    docx_path = base / "report.docx"
    pdf_path  = base / "report.pdf"
    _render_docx(tpl["title"], sections_out, docx_path)
    _render_pdf(tpl["title"],  sections_out, pdf_path)

    return {
        "status":       "success",
        "report_id":    report_id,
        "template":     template_id,
        "title":        tpl["title"],
        "role":         role,
        "generated_at": datetime.utcnow().isoformat() + "Z",
        "sections":     sections_out,
        "markdown":     full_md,
        "files": {
            "docx": f"/api/v1/reports/{report_id}/file/docx",
            "pdf":  f"/api/v1/reports/{report_id}/file/pdf",
            "md":   f"/api/v1/reports/{report_id}/file/md",
        },
    }


def get_report_file(report_id: str, fmt: str) -> Path:
    """Return the on-disk path to a generated report file."""
    if fmt not in {"docx", "pdf", "md"}:
        raise ValueError(f"Unsupported format: {fmt}")
    path = REPORTS_DIR / report_id / f"report.{fmt}"
    if not path.exists():
        raise FileNotFoundError(f"Report file not found: {report_id}.{fmt}")
    return path


# ═══════════════════════════════════════════════════════════════════════════
#   DATA-QUALITY SANITY CHECKS
# ═══════════════════════════════════════════════════════════════════════════

# Features that, if they appear as top predictors of revenue, are almost
# always leakage -- they're derived from or computed alongside the target.
_REVENUE_LEAKAGE_FEATURES = {
    "refund_amount", "is_returned", "return_reason",
    "cost_of_goods", "gross_margin_pct", "net_revenue",
    "discount_amount",
}


# ── Transparency-report helpers ─────────────────────────────────────────────
# These read the data_profile / labelling blocks attached by the v3.1 trainer
# (which in turn come from revenue_profiler, churn_labeler, growth_profiler).
# When a report is present, it represents a more-informed assessment than the
# heuristics below — the heuristics should respect it rather than override it.

_LEVEL_RANK = {"high": 3, "medium": 2, "low": 1}


def _get_report(kpis: dict, model_key: str) -> dict | None:
    """Return the transparency report for a model, or None."""
    block = kpis.get(model_key, {}) or {}
    # churn uses 'labelling', revenue/growth use 'data_profile'
    return block.get("data_profile") or block.get("labelling")


def _has_report_warning(report: dict | None, keyword: str) -> bool:
    """True if the report itself raised a warning containing `keyword`."""
    if not report:
        return False
    for w in report.get("warnings") or []:
        if keyword.lower() in str(w).lower():
            return True
    return False


def _floor_level(heuristic_level: str, report_level: str | None) -> str:
    """
    Combine a heuristic confidence level with the report's level.
    The result is the LOWER of the two — data-quality concerns can downgrade
    a strong-looking model, but a report's 'high' rating cannot rescue a
    model that performed poorly on its metrics.
    """
    if not report_level:
        return heuristic_level
    h = _LEVEL_RANK.get(heuristic_level, 2)
    r = _LEVEL_RANK.get(report_level,    2)
    final_rank = min(h, r)
    for name, rank in _LEVEL_RANK.items():
        if rank == final_rank:
            return name
    return heuristic_level


def _data_quality_warnings(kpis: dict) -> list[str]:
    """
    Inspect the evaluation results for statistically implausible numbers and
    return a list of plain-English warnings the user should see.
    """
    warnings: list[str] = []
    if not kpis:
        return warnings

    # ── Revenue model checks ──────────────────────────────────────────
    rev = kpis.get("revenue_model", {})
    if rev:
        try:
            r2 = float(rev.get("R2", 0))
            if r2 >= 0.95:
                warnings.append(
                    f"**Revenue model R² is {r2:.4f}** — values above 0.95 in "
                    "real-world business data almost always indicate target "
                    "leakage. Inspect the top features for any column that is "
                    "computed from or alongside revenue."
                )
        except (TypeError, ValueError):
            pass

        # Flat forecast detection
        fc = rev.get("12_period_forecast", [])
        if fc and len(fc) >= 2:
            mean = sum(fc) / len(fc)
            spread = (max(fc) - min(fc)) / mean if mean else 0
            if spread < 0.02:
                warnings.append(
                    f"**Revenue forecast is unusually flat** (range "
                    f"${max(fc) - min(fc):.2f} on a mean of ${mean:,.2f}). "
                    "The model may be predicting the mean rather than capturing "
                    "real time-series signal."
                )

        # Leakage-feature detection
        top_feats = {f.get("feature", "").lower() for f in rev.get("top_features", [])[:5]}
        leaks = top_feats & _REVENUE_LEAKAGE_FEATURES
        if leaks:
            warnings.append(
                f"**Likely target-leakage features in revenue model**: "
                f"{', '.join(sorted(leaks))}. These are typically computed from "
                "or alongside revenue, so the model may be solving an arithmetic "
                "identity rather than forecasting."
            )

    # ── Churn model checks ────────────────────────────────────────────
    churn = kpis.get("churn_model", {})
    if churn:
        labelling = churn.get("labelling")
        # If the labeller produced this rate using a high-confidence strategy
        # (forward_holdout) and didn't flag it, the rate is an OBSERVED business
        # fact about a defined cohort, not a labelling artefact. Skip the
        # heuristic flag — it would contradict the more-informed assessment.
        labeller_validated_rate = (
            labelling
            and labelling.get("confidence") == "high"
            and not _has_report_warning(labelling, "rate")
        )
        try:
            rate = float(churn.get("churn_rate_pct", 0))
            if rate >= 40 and not labeller_validated_rate:
                warnings.append(
                    f"**Churn rate of {rate:.1f}% is unusually high.** Real "
                    "ecommerce churn typically sits at 5–25% annually, SaaS at "
                    "2–8% monthly. Review the churn-label definition: a too-short "
                    "inactivity window will overcount seasonal buyers as churned."
                )
        except (TypeError, ValueError):
            pass

    # ── Growth model checks ───────────────────────────────────────────
    grow = kpis.get("growth_model", {})
    if grow:
        grow_profile = grow.get("data_profile")
        # The profiler computed a defensible horizon (safe_horizon_months) and
        # the trainer respected it. If both are present and agree, skip the
        # heuristic flag — it was designed for the old over-extrapolation bug
        # which the profiler now prevents.
        horizon = grow.get("forecast_horizon_months")
        safe = (grow_profile or {}).get("safe_horizon_months")
        profiler_certified = (
            grow_profile
            and horizon is not None
            and safe is not None
            and horizon <= safe
            and not _has_report_warning(grow_profile, "horizon")
        )
        train_months = grow.get("train_months", 0)
        try:
            train_months = int(train_months)
        except (TypeError, ValueError):
            train_months = 0
        forecast = grow.get("forecast") or grow.get("12_month_forecast", [])
        if (train_months and len(forecast) > train_months / 2
                and not profiler_certified):
            warnings.append(
                f"**Growth forecast horizon is large relative to training data** "
                f"({len(forecast)} months forecast from {train_months} months of "
                "history). Treat anything past the first 3–6 months as low-confidence."
            )

    # ── Cross-model consistency checks ────────────────────────────────
    # These catch contradictions between models that no single section
    # would notice on its own.
    rev_fc   = rev.get("12_period_forecast", []) if rev else []
    grow_fc  = (grow.get("forecast") or grow.get("12_month_forecast", [])) if grow else []

    if rev_fc and grow_fc:
        rev_mean   = sum(rev_fc) / len(rev_fc)
        grow_vals  = [
            float(x.get("forecast", 0)) if isinstance(x, dict) else float(x)
            for x in grow_fc
        ]
        grow_mean  = sum(grow_vals) / len(grow_vals) if grow_vals else 0

        # If the two means differ by more than 100x, units almost certainly
        # don't match (e.g. one is per-transaction, the other is per-month
        # aggregate). That's worth surfacing.
        if grow_mean and rev_mean:
            ratio = grow_mean / rev_mean
            if ratio > 100 or ratio < 0.01:
                warnings.append(
                    f"**Scale mismatch between revenue and growth models.** "
                    f"The revenue model forecasts **${rev_mean:,.2f} per "
                    f"transaction** (each prediction = one order); the growth "
                    f"model forecasts **${grow_mean:,.2f} per month** "
                    f"(each prediction = an entire month's revenue). The "
                    f"{ratio:,.0f}x gap is NOT a model error -- it's that "
                    "the two models answer different questions. Use the "
                    "revenue model for per-order pricing and discount "
                    "decisions; use the growth model for monthly budgeting "
                    "and capacity planning. To bridge them: the revenue "
                    "model's per-transaction average should reconcile "
                    f"approximately to the growth model when multiplied by "
                    f"your monthly order count "
                    f"(e.g. ${rev_mean:,.0f} per order x ~"
                    f"{int(grow_mean / rev_mean):,} orders/month "
                    f"≈ ${grow_mean:,.0f}/month). If that reconciliation "
                    "doesn't hold, one of the models needs investigation."
                )

    # MAPE sanity -- if the per-transaction revenue model has a MAPE that
    # would be alarming on aggregated data but is normal for transactional
    # data, surface that distinction so the report doesn't over-react.
    # If the profiler already raised the equivalent CoV warning, skip this
    # one to avoid duplicating the same finding in different language.
    rev_profile          = (rev or {}).get("data_profile")
    profiler_cov_flagged = _has_report_warning(rev_profile, "coefficient of variation")
    try:
        rev_mape = float(rev.get("MAPE_pct", 0)) if rev else 0
        rev_r2   = float(rev.get("R2", 0))       if rev else 0
        if rev_mape >= 30 and rev_r2 >= 0.7 and not profiler_cov_flagged:
            warnings.append(
                f"**Revenue MAPE is {rev_mape:.1f}% but R² is {rev_r2:.3f}.** "
                "On per-transaction data with mixed order sizes, high MAPE is "
                "expected (small orders inflate the metric). Combined with a "
                f"healthy R² of {rev_r2:.3f}, the model is fitting well in "
                "absolute terms -- treat MAE and sMAPE as the primary error "
                "metrics, not MAPE."
            )
    except (TypeError, ValueError):
        pass

    return warnings


# ═══════════════════════════════════════════════════════════════════════════
#   MODEL CONFIDENCE LAYER
# ═══════════════════════════════════════════════════════════════════════════
#
# Programmatic per-model confidence ratings, computed deterministically
# from training size, metric quality, and detected warnings. These ratings
# are surfaced TWICE: once in the report's opening "Data Quality &
# Confidence" section, and once inside every LLM prompt so the model's
# language matches its actual reliability.
# ───────────────────────────────────────────────────────────────────────────

def _model_confidence(kpis: dict, warnings: list[str]) -> dict:
    """
    Return a dict like:
        {
          "revenue": {
              "level": "medium",
              "reasons": ["..."],
              "use_for": "Aggregate revenue planning over weeks/months.",
              "do_not_use_for": "Per-order pricing decisions or quoting "
                                "individual transaction values.",
          },
          ...
        }
    """
    out: dict = {}
    warning_blob = " ".join(warnings).lower() if warnings else ""

    # ── Revenue model ─────────────────────────────────────────────────
    rev = kpis.get("revenue_model", {})
    if rev:
        reasons: list[str] = []
        use_for = ""
        do_not  = ""
        try:
            r2    = float(rev.get("R2", 0))
            mape  = float(rev.get("MAPE_pct", 0))
            train = int(rev.get("train_rows", 0))
        except (TypeError, ValueError):
            r2, mape, train = 0, 0, 0

        # Confidence is gated on three things: training size, R² range,
        # AND MAPE. A model can have a great R² but lousy per-row error
        # (common on transactional data with mixed order sizes), in which
        # case it's reliable on AVERAGES but noisy on individual rows.
        # That kind of model deserves Medium, not High.
        level = "medium"
        if r2 >= 0.95:
            # Implausible R² -- almost always leakage. See warnings.
            level = "low"
            reasons.append(
                f"R² of {r2:.4f} is suspiciously high -- possible target leakage. "
                "Treat all revenue numbers as provisional until features are re-audited."
            )
            use_for = "Nothing until the leakage suspicion is resolved."
            do_not  = ("Any revenue forecasting, target-setting, or pricing "
                       "decision -- numbers are likely a near-tautology.")
        elif train < 5000 or r2 < 0.4:
            level = "low"
            reasons.append(
                f"Only {train:,} training rows / R² of {r2:.3f}. Forecasts are "
                "directional at best."
            )
            use_for = "Rough trend estimation only."
            do_not  = "Any decision involving real budget or staff time."
        elif (train >= 20000 and 0.65 <= r2 <= 0.94
              and mape < 30 and "leakage" not in warning_blob):
            level = "high"
            reasons.append(
                f"Trained on {train:,} rows with R² of {r2:.3f} and MAPE of "
                f"{mape:.1f}%. Reliable in aggregate AND on individual rows."
            )
            use_for = ("Both aggregate revenue planning AND per-transaction "
                       "predictions (e.g. dynamic pricing).")
            do_not  = ("Decisions that require causal interpretation -- "
                       "feature importance is correlation, not causation.")
        elif (train >= 20000 and 0.65 <= r2 <= 0.94
              and "leakage" not in warning_blob):
            # Big training set + healthy R² but high per-row MAPE -> Medium.
            level = "medium"
            reasons.append(
                f"Trained on {train:,} rows with R² of {r2:.3f} -- aggregate "
                f"forecasts are reliable. However, MAPE of {mape:.1f}% means "
                "per-transaction predictions are noisy: small orders carry "
                "large percentage errors."
            )
            use_for = ("Aggregate revenue planning over weeks or months "
                       "(totals, averages, trends).")
            do_not  = ("Per-order pricing, quoting individual transaction "
                       "values, or any decision that hinges on a single "
                       "row's predicted revenue.")
        else:
            reasons.append(
                f"Trained on {train:,} rows, R² of {r2:.3f}, MAPE {mape:.1f}%. "
                "Acceptable for aggregate planning, not for per-row decisions."
            )
            use_for = "High-level aggregate planning."
            do_not  = "Per-row decisions or causal explanations."
        out["revenue"] = {
            "level":          level,
            "reasons":        reasons,
            "use_for":        use_for,
            "do_not_use_for": do_not,
        }

    # ── Churn model ───────────────────────────────────────────────────
    churn = kpis.get("churn_model", {})
    if churn:
        reasons = []
        try:
            auc   = float(churn.get("ROC_AUC", 0))
            rate  = float(churn.get("churn_rate_pct", 0))
            train = int(churn.get("train_rows", 0))
        except (TypeError, ValueError):
            auc, rate, train = 0, 0, 0

        # Aggressive churn-label rate is the dominant signal here -- it
        # makes EVERY downstream metric suspect, regardless of AUC.
        if rate >= 40:
            level = "low"
            reasons.append(
                f"Churn rate of {rate:.1f}% is implausibly high; the churn-label "
                "rule almost certainly overcounts seasonal buyers. Even an "
                f"AUC of {auc:.3f} is not enough to redeem this -- the model "
                "is sorting on a flawed target. Resolve this before acting on "
                "any churn-related recommendation."
            )
            use_for = ("Sanity-checking the churn-label rule itself "
                       "(the high rate IS the diagnostic).")
            do_not  = ("Customer-level retention targeting, retention budget "
                       "allocation, or any segmentation that uses the model's "
                       "risk scores. Output is unreliable until the label "
                       "rule is fixed.")
        elif auc >= 0.80 and train >= 5000:
            level = "high"
            reasons.append(
                f"AUC of {auc:.3f} on {train:,} training customers; strong "
                "discriminator with a plausible churn rate."
            )
            use_for = ("Ranking customers by churn risk for retention "
                       "campaigns. Top-decile prioritisation is reliable.")
            do_not  = ("Treating individual probability scores as exact "
                       "(e.g. '73% chance of churn') -- AUC measures ranking, "
                       "not calibration.")
        elif auc >= 0.70:
            level = "medium"
            reasons.append(
                f"AUC of {auc:.3f} is decent. Treat individual risk scores as "
                "directional rather than precise."
            )
            use_for = ("Broad high-risk vs low-risk segmentation for "
                       "experimentation.")
            do_not  = ("Fine-grained risk rankings or claims about exact "
                       "churn probabilities.")
        else:
            level = "low"
            reasons.append(
                f"AUC of {auc:.3f} is weak -- the model barely outperforms "
                "random for ranking customers by churn risk."
            )
            use_for = "Hypothesis generation only; treat output as exploratory."
            do_not  = ("Any production retention decision -- ranking is too "
                       "close to random.")
        out["churn"] = {
            "level":          level,
            "reasons":        reasons,
            "use_for":        use_for,
            "do_not_use_for": do_not,
        }

    # ── Growth model ──────────────────────────────────────────────────
    grow = kpis.get("growth_model", {})
    if grow:
        reasons = []
        try:
            r2    = float(grow.get("R2", 0))
            train_m = int(grow.get("train_months", 0))
            test_m  = int(grow.get("test_months", 0))
        except (TypeError, ValueError):
            r2, train_m, test_m = 0, 0, 0

        # Time-series models with small training windows are inherently
        # low-confidence regardless of in-sample fit. R² on 4 test months
        # is not the same as 12 -- statistical power matters, and a single
        # test season can hide structural drift.
        if train_m < 24 or test_m < 6:
            level = "low"
            reasons.append(
                f"Trained on only {train_m} months and tested on {test_m} -- "
                "less than one full seasonal cycle. A high test-set R² on "
                f"this little data ({r2:.3f}) is not statistically meaningful: "
                "it could equally reflect the model picking up signal OR "
                "luck on a small holdout."
            )
            use_for = ("Directional sense of trajectory (up / flat / down) "
                       "for next 1-3 months only.")
            do_not  = ("Budget commitments, hiring decisions, supplier "
                       "contracts, or any specific monthly figure as a "
                       "target. Re-evaluate after 24+ months of training "
                       "data and 6+ test months exist.")
        elif r2 >= 0.7:
            level = "medium"
            reasons.append(
                f"Trained on {train_m} months, R² of {r2:.3f}. Reasonable "
                "for short-horizon planning (3-6 months out); not for "
                "year-end commitments."
            )
            use_for = ("Quarterly capacity and budget planning over a "
                       "3-6 month horizon.")
            do_not  = ("Year-end commitments or any planning horizon "
                       "longer than 6 months.")
        else:
            level = "low"
            reasons.append(
                f"R² of {r2:.3f} on {train_m} months is weak. Treat the "
                "12-month forecast as a placeholder."
            )
            use_for = "Trend direction only."
            do_not  = "Any specific monthly figure for planning."
        out["growth"] = {
            "level":          level,
            "reasons":        reasons,
            "use_for":        use_for,
            "do_not_use_for": do_not,
        }

    # ── Report-driven post-pass ───────────────────────────────────────
    # Where a transparency report exists for a model (data_profile or
    # labelling), adjust the heuristic result so the more-informed
    # assessment wins:
    #   - level: take the LOWER of (heuristic level, report level)
    #            so a high-confidence labeller can't rescue a model with
    #            poor metrics, but a strong-metrics model whose data the
    #            profiler flagged still gets downgraded.
    #   - SPECIAL CASE: when the labeller is high-confidence and the
    #            heuristic only said "low" because of the rate (e.g.
    #            50% churn from forward-holdout), restore the level
    #            and rephrase the reasoning. The rate is a real
    #            business fact, not a labelling artefact.
    #   - reasons: prepend the report's `definition` so the LLM sees
    #            HOW the metric was computed before reading WHY we
    #            scored its confidence the way we did.
    #   - use_for / do_not_use_for: prefer the report's authoritative
    #            suitable_for / not_suitable_for lists where present.
    for model_key in ("revenue", "churn", "growth"):
        if model_key not in out:
            continue
        # Map UI model_key → kpis key
        kpis_key = {
            "revenue": "revenue_model",
            "churn":   "churn_model",
            "growth":  "growth_model",
        }[model_key]
        report = _get_report(kpis, kpis_key)
        if not report:
            continue

        existing       = out[model_key]
        report_level   = report.get("confidence")
        report_reasons = []

        # Special case for churn: a high-confidence labeller overrides the
        # rate-based "low confidence" downgrade, since the rate is observed.
        if (model_key == "churn"
                and report_level == "high"
                and existing["level"] == "low"
                and not _has_report_warning(report, "rate")):
            churn = kpis.get("churn_model", {})
            try:
                auc   = float(churn.get("ROC_AUC", 0))
                rate  = float(churn.get("churn_rate_pct", 0))
                train = int(churn.get("train_rows", 0))
            except (TypeError, ValueError):
                auc, rate, train = 0, 0, 0
            # Recompute as if the rate-rule didn't fire
            if auc >= 0.80 and train >= 5000:
                existing["level"] = "high"
                existing["reasons"] = [
                    f"AUC of {auc:.3f} on {train:,} training customers; "
                    f"strong discriminator. Churn rate of {rate:.1f}% is an "
                    "observed business outcome on a labeller-validated cohort, "
                    "not a measurement artefact."
                ]
                existing["use_for"] = (
                    "Ranking customers by churn risk for retention campaigns. "
                    "Top-decile prioritisation is reliable."
                )
                existing["do_not_use_for"] = (
                    "Treating individual probability scores as exact "
                    "(e.g. '73% chance of churn') -- AUC measures ranking, "
                    "not calibration."
                )
            elif auc >= 0.70:
                existing["level"]   = "medium"
                existing["reasons"] = [
                    f"AUC of {auc:.3f} on a labeller-validated cohort. "
                    "Treat individual risk scores as directional."
                ]
                report_reasons = []  # already covered above

        # General case: take the lower of heuristic level and report level
        existing["level"] = _floor_level(existing["level"], report_level)

        # Prepend the report's own definition so the LLM sees the provenance
        if report.get("definition"):
            report_reasons.append(report["definition"])

        if report_reasons:
            existing["reasons"] = report_reasons + (existing.get("reasons") or [])

        # Prefer the report's curated suitable_for / not_suitable_for lists
        if report.get("suitable_for"):
            existing["use_for"] = "; ".join(report["suitable_for"])
        if report.get("not_suitable_for"):
            existing["do_not_use_for"] = "; ".join(report["not_suitable_for"])

    return out


_CONF_BADGE = {
    "high":   "🟢 High",
    "medium": "🟡 Medium",
    "low":    "🔴 Low",
}


def _format_confidence_md(confidence: dict) -> str:
    """Render the confidence dict as a markdown table for the report opener."""
    if not confidence:
        return ""
    lines = ["**Model Confidence Summary**", "",
             "| Model | Confidence | Why | ✅ Use For | ❌ Do NOT Use For |",
             "|---|---|---|---|---|"]
    pretty = {"revenue": "Revenue Forecast", "churn": "Customer Churn",
              "growth": "Growth Outlook"}
    for key, info in confidence.items():
        badge   = _CONF_BADGE.get(info["level"], info["level"])
        why     = " ".join(info.get("reasons", [])).replace("\n", " ").replace("|", "/")
        use_for = (info.get("use_for", "") or "—").replace("\n", " ").replace("|", "/")
        do_not  = (info.get("do_not_use_for", "") or "—").replace("\n", " ").replace("|", "/")
        lines.append(f"| {pretty.get(key, key)} | {badge} | {why} | {use_for} | {do_not} |")
    return "\n".join(lines)


def _format_confidence_compact(confidence: dict) -> str:
    """Compact rendering for inclusion in LLM prompts. Includes
    use-for and do-not-use-for so the LLM knows the boundaries of
    each model's reliable application."""
    if not confidence:
        return ""
    pretty = {"revenue": "Revenue", "churn": "Churn", "growth": "Growth"}
    out = []
    for k, info in confidence.items():
        out.append(
            f"- {pretty.get(k, k)}: {info['level'].upper()} confidence -- "
            f"{' '.join(info.get('reasons', []))}"
        )
        if info.get("use_for"):
            out.append(f"    USE FOR: {info['use_for']}")
        if info.get("do_not_use_for"):
            out.append(f"    DO NOT USE FOR: {info['do_not_use_for']}")
    return "\n".join(out)


# ── Decision-gate matrix ───────────────────────────────────────────────
# Decision gating is a 2-D function of CONFIDENCE x IMPACT, not just
# confidence. A low-confidence growth model is a planning headache; a
# low-confidence churn model is a brand-damage risk. Both deserve "Block",
# but only the latter deserves urgent human review.
#
# Impact is per-decision, not per-model -- it reflects how much business
# pain a wrong call creates. We hard-code it here because the model itself
# can't know.
_DECISION_IMPACT = {
    # Wrong pricing call: lost margin, but recoverable next cycle.
    "revenue": "medium",
    # Wrong churn call: either alienate loyal customers with retention spam
    # or under-spend on at-risk customers and lose them. Trust + revenue
    # damage, hard to recover from.
    "churn":   "high",
    # Wrong growth call: over/under-budget capacity, shows up over a
    # quarter, recoverable.
    "growth":  "medium",
}

# Per-decision business label (left column of the gate table)
_DECISION_LABEL = {
    "revenue": "Pricing & revenue optimisation",
    "churn":   "Customer retention campaigns",
    "growth":  "Growth & capacity planning",
}

# 2-D gate matrix: gate_for_level_and_impact[confidence][impact]
# yields (symbol, short-text-rationale).
_GATE_MATRIX = {
    "high": {
        "high":   ("✅ Proceed",     "High confidence, well-validated."),
        "medium": ("✅ Proceed",     "High confidence, well-validated."),
        "low":    ("✅ Proceed",     "High confidence, well-validated."),
    },
    "medium": {
        "high":   ("⚠️ Pilot only",  "Medium confidence, high impact -- validate at small scale before full rollout."),
        "medium": ("⚠️ Pilot only",  "Validate at small scale before full rollout."),
        "low":    ("⚠️ Pilot only",  "Acceptable for a pilot."),
    },
    "low": {
        "high":   ("🚨 Escalate",    "Low confidence on a HIGH-impact decision -- requires human review before any commitment."),
        "medium": ("❌ Block",       "Address the underlying limitation before acting."),
        "low":    ("❌ Block",       "Address the underlying limitation before acting."),
    },
}


def _gate_for(level: str, impact: str) -> tuple[str, str]:
    """Return (symbol, rationale) for a given confidence x impact pair."""
    return _GATE_MATRIX.get(level, {}).get(impact) or ("?", "Unknown")


# Per-recommendation lookup, used by the synthesis prompt instructions.
# Same mapping as _GATE_MATRIX but flattened for easier prompt embedding.
_GATE_FOR_LEVEL = {
    "high":   ("✅ Proceed",     "Safe to act on."),
    "medium": ("⚠️ Pilot only",  "Validate with a small-scale pilot first."),
    "low":    ("❌ Block",       "Resolve the underlying issue before acting."),
}


_IMPACT_BADGE = {
    "high":   "🔥 High",
    "medium": "📊 Medium",
    "low":    "📉 Low",
}


def _format_decision_gates_md(confidence: dict) -> str:
    """
    Render the Decision Gate summary table -- the at-a-glance scorecard
    that tells leadership which decisions are safe to act on, which need
    piloting, which are blocked, and which need to be escalated for human
    review (high impact + low confidence).
    """
    if not confidence:
        return ""
    lines = [
        "**Decision Gate Summary**",
        "",
        "_Quick-glance verdict for each business decision the models inform. "
        "Gates combine model **confidence** with **impact** -- a wrong call "
        "on a high-impact decision needs human review even when the model is "
        "uncertain. Use this scorecard as your gate before committing budget "
        "or staff time._",
        "",
        "| Decision | Impact | Confidence | Gate | Rationale |",
        "|---|---|---|---|---|",
    ]
    for key, info in confidence.items():
        level    = info["level"]
        impact   = _DECISION_IMPACT.get(key, "medium")
        conf_b   = _CONF_BADGE.get(level, level)
        imp_b    = _IMPACT_BADGE.get(impact, impact)
        symbol, gate_rationale = _gate_for(level, impact)
        label    = _DECISION_LABEL.get(key, key.title())

        # Combine the gate's structural rationale with the model's specific
        # reason. Take up to the first 2 sentences (cap 300 chars) so we
        # capture both the report's definition AND any heuristic reasoning
        # behind a downgrade — single-sentence truncation was leaving cells
        # ending mid-thought and hiding the actual reason for low gates.
        full = " ".join(info.get("reasons", [])).strip()
        parts = full.split(". ", 2)
        if len(parts) >= 2:
            first = ". ".join(parts[:2]).strip()
        else:
            first = full.strip()
        if first and not first.endswith("."):
            first += "."
        if len(first) > 300:
            first = first[:297].rstrip() + "..."
        why = (gate_rationale + " " + first).replace("|", "/").replace("\n", " ").strip()
        lines.append(f"| {label} | {imp_b} | {conf_b} | {symbol} | {why} |")
    return "\n".join(lines)


def _ml_context_for(which: Optional[str], kpis: dict) -> str:
    """Format the requested ML model results into prompt-ready text."""
    if not which or not kpis:
        return ""
    parts: list[str] = []
    if which in ("revenue", "all"):
        s = _format_revenue_ml(kpis)
        if s: parts.append(s)
    if which in ("churn", "all"):
        s = _format_churn_ml(kpis)
        if s: parts.append(s)
    if which in ("growth", "all"):
        s = _format_growth_ml(kpis)
        if s: parts.append(s)
    return "\n\n".join(parts)


def _rag_context_for(query: str, n_results: int = 3) -> str:
    """Pull the top-N most relevant uploaded-data chunks for a topical query."""
    try:
        result = query_knowledge_base(query, category=None, n_results=n_results)
        return result.get("context", "") if result.get("count", 0) > 0 else ""
    except Exception:
        return ""


def _build_section_prompt(
    section_title: str,
    instructions: str,
    role: str,
    ml_context: str,
    rag_context: str,
    confidence: Optional[dict] = None,
) -> str:
    pieces = [
        f"You are a senior {role} writing one section of a formal business report.",
        f"SECTION TITLE: {section_title}",
        "",
        # Hard anti-reasoning instruction. Without this some free-tier models
        # (DeepSeek-V3 in particular) write their entire chain-of-thought into
        # the section content -- "We need to write...", "Let me think...",
        # "Sentence 1: ..." -- and run out of token budget before producing
        # a finished answer.
        "OUTPUT CONTRACT:",
        "- Output ONLY the finished section prose. Nothing else.",
        "- Do NOT show your reasoning, planning, or working-out.",
        "- Do NOT begin with phrases like 'We need to', 'Let me', 'I need to', "
        "'Let's', 'First, I'll', 'Sentence 1:', 'Step 1:', 'Thinking:', or "
        "'The user wants'. Begin DIRECTLY with the first word of the actual "
        "section.",
        "- Do NOT label your own output ('Here is the section:', 'Final answer:'). "
        "Just write the section.",
        "",
        "INSTRUCTIONS:",
        instructions,
        "",
        "STYLE:",
        "- Plain business prose. No markdown headings (the section header is added separately).",
        "- Concrete numbers wherever the data supports them.",
        "- No filler. No hedging language like 'it should be noted that'.",
        "- Do NOT invent facts. If a number isn't in the data below, do not state it.",
        "- Do NOT invent target numbers in any forward-looking statement. "
        "Phrases like 'reduce churn by 10% in 90 days' are forbidden unless "
        "the 10% figure comes directly from the model output. If you cannot "
        "anchor a target to the data, say 'reduce churn' without the percentage.",
        "",
        "CONFIDENCE-AWARE LANGUAGE:",
        "Each model has a confidence rating (high / medium / low) shown below. "
        "Match your language to the rating:",
        "- HIGH confidence: assert findings directly. 'The model predicts X.'",
        "- MEDIUM confidence: hedge appropriately. 'The model suggests X, "
        "with material uncertainty given the limited training history.'",
        "- LOW confidence: explicitly caveat. 'The model produces a figure "
        "of X, but should be treated as directional only -- training data "
        "is too limited to support firm conclusions.'",
        "Never present a low-confidence number with high-confidence framing.",
        "",
        "CRITICAL REASONING:",
        "Treat the model metrics with healthy scepticism, but do not invent "
        "concerns where none exist. In particular:",
        "- An R-squared of 0.85 to 0.94 is NORMAL for well-tuned business "
            "models -- DO NOT flag it as leakage. Only flag R-squared >= 0.95 "
            "as a possible leakage signal, and only when paired with "
            "implausible MAPE.",
        "- A MAPE above 30% on per-transaction data with mixed order sizes is "
            "common and is NOT a model failure -- small orders inflate MAPE. "
            "Report MAPE but do not treat it alone as evidence the model is "
            "unreliable. Lean on sMAPE, MAE, and R-squared together.",
        "- A forecast is 'unusually flat' only if its range is below 2% of "
            "the mean. Do not call a 5%+ range flat.",
        "- A churn rate above 40% is genuinely unusual and SHOULD be flagged. "
            "Below 40%, do not raise it as a concern.",
        "- A feature is leakage only if it is computed FROM the target value "
            "(e.g. refund_amount, cost_of_goods, gross_margin_pct, is_returned "
            "for a revenue model). Do NOT speculate that any high-importance "
            "feature 'might be leakage' -- only call it out by name when it "
            "matches the known leakage list.",
        "- Avoid superlatives ('strong', 'excellent', 'poised for significant "
            "revenue'). Prefer measured language.",
        "- Feature importance is correlation, not causation. Do NOT write "
            "'increasing X will boost Y', 'X causes Y', or 'focusing on X "
            "will lift Y'. Use 'X is the strongest predictor of Y', "
            "'X is associated with higher Y', 'X is correlated with Y'.",
        "- If two model results contradict each other (e.g. revenue per period "
            "of $691 from one model vs. $1.2M per month from another), flag "
            "the contradiction explicitly rather than glossing over it.",
        "- If you see TWO different forecasts for the same model (e.g. one "
            "snapshot says $843, another says $691), DO NOT treat them as "
            "rival models. Use the single set of ML MODEL RESULTS provided "
            "below as the source of truth and ignore any older snapshots "
            "that may appear in the uploaded-data context.",
        "",
        "FORMATTING:",
        "- DO NOT repeat the section title at the start of your output. "
        "The section heading is added separately, so beginning with "
        "'### Revenue Performance' or even just 'Revenue Performance' "
        "creates duplicate headings in the rendered report.",
        "- Open the section with a substantive sentence, not a heading.",
        "",
    ]
    if confidence:
        pieces += [
            "-- MODEL CONFIDENCE RATINGS --------------------------------",
            _format_confidence_compact(confidence),
            "",
        ]
    if ml_context:
        pieces += ["-- ML MODEL RESULTS (source of truth) ---------------------",
                   ml_context, ""]
    if rag_context:
        pieces += ["-- UPLOADED DATA CONTEXT -----------------------------------",
                   rag_context, ""]
    pieces.append("Write the section now:")
    return "\n".join(pieces)


def _build_synthesis_prompt(
    section_title: str,
    instructions: str,
    role: str,
    ml_context: str,
    rag_context: str,
    prior_draft: str,
    warnings: list[str],
    confidence: Optional[dict] = None,
) -> str:
    """
    Build the prompt for a synthesis section (recommendations / priority
    actions). Differs from _build_section_prompt by also feeding the LLM
    the FULL DRAFT of the rest of the report and the data-quality warnings,
    so its actions reference the actual analysis -- not just raw metrics.
    """
    pieces = [
        f"You are a senior {role} writing the final synthesis section of a "
        "business report. The rest of the report has already been written; "
        "it is provided below as REPORT DRAFT.",
        f"SECTION TITLE: {section_title}",
        "",
        # Same anti-reasoning gate as _build_section_prompt. Priority Actions
        # is the section most prone to chain-of-thought leakage because it
        # involves multi-step planning -- this is where it matters most.
        "OUTPUT CONTRACT:",
        "- Output ONLY the finished section. Nothing else.",
        "- Do NOT show your reasoning, planning, or working-out.",
        "- Do NOT begin with 'We need to', 'Let me', 'Let's', 'First, I'll', "
        "'Step 1:', 'Thinking:', or any meta-narrative.",
        "- Begin directly with the first numbered action.",
        "- Do NOT label your own output ('Here are the recommendations:'). "
        "Just write them.",
        "",
        "INSTRUCTIONS:",
        instructions,
        "",
        "ADDITIONAL RULES FOR SYNTHESIS SECTIONS:",
        "- Each recommendation MUST reference a specific finding from the "
            "REPORT DRAFT (e.g. 'given the 65% churn rate flagged in the "
            "Customer Churn Risk section') -- not just a raw metric.",
        "- Recommendations should be BUSINESS actions, not 'fix the model' "
            "actions, UNLESS the data-quality warnings make a model fix the "
            "highest-priority item. Examples of business actions: launch "
            "retention campaign for X segment, audit pricing on Y category, "
            "negotiate supplier terms for top-N SKUs.",
        "- If two sections of the report contradict each other or the data-"
            "quality warnings raise a critical concern, the FIRST recommendation "
            "must be to resolve that concern -- everything else is downstream.",
        "- No vague language. 'Refine the model' is not actionable; 'Audit "
            "the churn-label rule and re-run training' is.",
        "- Anchor each action to a specific number from the actual analysis.",
        "",
        "CONFIDENCE TAGGING (mandatory):",
        "Append a confidence tag to EVERY recommendation, in the format "
        "`[Confidence: high|medium|low]`. The confidence is determined by "
        "the model the recommendation is anchored to (see model confidence "
        "ratings below). Recommendations that depend on a flagged data-"
        "quality warning are automatically LOW confidence until the warning "
        "is resolved.",
        "",
        "DECISION GATE (mandatory):",
        "After the confidence tag, append a decision gate symbol. The gate "
        "is determined by a 2-D matrix of CONFIDENCE x IMPACT, not just "
        "confidence. Impact reflects how much business pain a wrong call "
        "creates:",
        "  - Pricing / revenue decisions  = MEDIUM impact",
        "  - Customer retention / churn   = HIGH impact "
            "(misclassifying loyal customers damages trust)",
        "  - Growth / capacity planning   = MEDIUM impact",
        "",
        "Pick the gate using this matrix:",
        "  | Confidence | Low impact | Medium impact | High impact |",
        "  |------------|-----------|---------------|-------------|",
        "  | HIGH       | ✅ Proceed | ✅ Proceed    | ✅ Proceed  |",
        "  | MEDIUM     | ⚠️ Pilot  | ⚠️ Pilot     | ⚠️ Pilot   |",
        "  | LOW        | ❌ Block  | ❌ Block      | 🚨 Escalate |",
        "",
        "🚨 Escalate is the critical case: low confidence + high impact "
        "means the cost of getting it wrong is too high to either block "
        "outright (which leaves money on the table) or proceed (which "
        "risks brand damage). It explicitly demands human review.",
        "",
        "Format each recommendation like this exactly (one line summary "
        "after the action description):",
        "  > `[Confidence: low] 🚨 Escalate — high-impact churn decision "
        "with low-quality labels; bring to leadership review before any "
        "retention budget allocation.`",
        "",
        "If a recommendation is itself the action to RESOLVE a low-confidence "
        "warning (e.g. 'audit the churn-label rule'), it is exempt from the "
        "Block / Escalate gate and should be marked `✅ Proceed` because "
        "resolving the warning is always safe to do.",
        "",
        "TARGET-NUMBER PROHIBITION:",
        "Do NOT invent target numbers. Forbidden: 'reduce churn by 10%', "
        "'5% revenue increase in 90 days', 'lift retention by 15 points'. "
        "Allowed: 'reduce churn in the high-risk segment (currently 65%)', "
        "'improve retention among the lowest-engagement customers'. If "
        "you cannot anchor a target to an actual model output, omit the "
        "target entirely.",
        "",
        "SEGMENT-CUTOFF PROHIBITION:",
        "Do NOT invent customer segment cutoffs that the model didn't "
        "produce. Forbidden: 'customers with fewer than 2 orders', "
        "'customers with avg_days_between_orders > 30', 'customers with "
        "predicted_clv below $X'. The model gives you FEATURE IMPORTANCE, "
        "not THRESHOLDS. Without an explicit threshold from the model "
        "output, refer to the segment by feature name only: 'customers "
        "in the lowest quartile of total_orders', 'customers ranked "
        "high-risk by the model', 'segment defined by the top churn "
        "predictor (total_orders)'. The data team can attach a numeric "
        "cutoff later when targeting.",
        "",
        "CAUSAL-LANGUAGE PROHIBITION:",
        "Feature importance is correlation, not causation. Forbidden: "
        "'increasing order frequency will boost revenue', 'customers "
        "churn BECAUSE they have low order counts', 'discounts CAUSE "
        "revenue lift'. Allowed: 'order frequency is correlated with "
        "revenue', 'low order count is the strongest predictor of "
        "churn', 'discounted price is associated with higher revenue'. "
        "If you want to make a causal claim, qualify it with 'subject "
        "to A/B test validation'.",
        "",
        "TWO-SNAPSHOT WARNING:",
        "If the REPORT DRAFT or RAG context mentions multiple revenue "
        "forecast values that disagree (e.g. $843 vs $691), DO NOT treat "
        "them as rival models. They are different snapshots of the same "
        "model from different training runs. The current ML MODEL RESULTS "
        "(below) are the source of truth. If you spot this in the draft, "
        "your first recommendation should be to clear the stale snapshot "
        "from the knowledge base.",
        "",
    ]
    if confidence:
        pieces += [
            "-- MODEL CONFIDENCE RATINGS --------------------------------",
            _format_confidence_compact(confidence),
            "",
        ]
    if ml_context:
        pieces += ["-- ML MODEL RESULTS (source of truth) ---------------------",
                   ml_context, ""]
    if rag_context:
        pieces += ["-- UPLOADED DATA CONTEXT -----------------------------------",
                   rag_context, ""]
    if warnings:
        pieces += ["-- DATA-QUALITY WARNINGS ALREADY SURFACED -----------------",
                   "\n".join(f"- {w}" for w in warnings), ""]
    if prior_draft:
        pieces += ["-- REPORT DRAFT (previously generated sections) -----------",
                   prior_draft, ""]
    pieces.append("Write the synthesis section now:")
    return "\n".join(pieces)


# ═══════════════════════════════════════════════════════════════════════════
#   RENDERING
# ═══════════════════════════════════════════════════════════════════════════

def _strip_duplicate_heading(content: str, section_title: str) -> str:
    """
    Strip a duplicate section title that the LLM sometimes prepends to its
    output (e.g. '### Customer Churn Risk' followed by the actual prose).
    The renderer adds the heading separately, so leaving it in causes the
    title to appear twice in the rendered report.

    Tolerates leading whitespace, '#' / '##' / '###' / '####' prefixes, and
    bold-wrapped variants like '**Customer Churn Risk**'.
    """
    if not content or not section_title:
        return content
    title_l = section_title.strip().lower()
    lines   = content.splitlines()
    # Find the first non-empty line
    for i, line in enumerate(lines):
        if not line.strip():
            continue
        # Strip markdown heading markers + bold
        stripped = line.strip()
        cleaned  = stripped.lstrip("#").strip()
        cleaned  = cleaned.strip("*").strip()
        cleaned  = cleaned.rstrip(":").strip()
        if cleaned.lower() == title_l:
            # Drop this line and any blank line after it
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            return "\n".join(lines[j:])
        # First non-empty line is not the title -- nothing to strip
        return content
    return content


def _stitch_markdown(title: str, sections: list[dict]) -> str:
    today = datetime.utcnow().strftime("%B %d, %Y")
    lines = [
        f"# {title}",
        f"_Generated by OptimaAi on {today}_",
        "",
    ]
    for sec in sections:
        lines += [f"## {sec['title']}", "", sec["content_md"], ""]
    return "\n".join(lines)


def _render_docx(title: str, sections: list[dict], out_path: Path) -> None:
    """Render a polished Word document using python-docx."""
    try:
        from docx import Document
        from docx.shared  import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        # Soft fallback -- don't break the whole report if python-docx isn't installed
        print("  [report] python-docx not installed; writing plain-text .docx fallback")
        out_path.write_text(_stitch_markdown(title, sections), encoding="utf-8")
        return

    doc = Document()

    # ── Default font ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title ──
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)

    # ── Subtitle ──
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        f"OptimaAi  ·  Generated {datetime.utcnow().strftime('%B %d, %Y')}"
    )
    sub_run.italic = True
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()  # spacer

    # ── Sections ──
    for sec in sections:
        h = doc.add_paragraph()
        h_run = h.add_run(sec["title"])
        h_run.bold = True
        h_run.font.size = Pt(14)
        h_run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)

        # Body -- split on double newline to preserve paragraphs
        body = sec["content_md"].strip()
        for para in body.split("\n\n"):
            para = para.strip()
            if not para:
                continue
            p = doc.add_paragraph(para)
            p.paragraph_format.space_after = Pt(6)

        doc.add_paragraph()  # spacer

    # ── Footer ──
    section = doc.sections[0]
    footer  = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("OptimaAi  ·  Confidential")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    doc.save(out_path)


# Emoji range strip — reportlab's default Helvetica font has no glyphs for
# the emoji characters used in confidence/gate badges, so they render as ■
# boxes in the PDF. Stripping them is lossless for content because every
# badge already includes a text label next to its emoji ("🟢 High" → "High").
import re as _re
_EMOJI_RE = _re.compile(
    "["
    "\U0001F300-\U0001FAFF"   # broad pictograph + emoticons + transport
    "\U00002600-\U000027BF"   # miscellaneous symbols + dingbats (✅❌⚠)
    "\U00002B00-\U00002BFF"   # misc symbols and arrows
    "\U0001F000-\U0001F0FF"
    "\U0000FE0F"              # variation selector-16
    "\U0000200D"              # zero-width joiner
    "]+",
    flags=_re.UNICODE,
)
_BOLD_MD_RE = _re.compile(r"\*\*([^*\n]+?)\*\*")


def _pdf_clean(text: str) -> str:
    """
    Prepare a markdown string for reportlab Paragraph rendering:
      1. Strip emoji (reportlab default fonts can't render them)
      2. Convert **bold** markdown to <b>bold</b> (reportlab's tag syntax)
      3. Collapse the double-spaces left behind by emoji removal
    The .docx and .md outputs are unaffected — they keep emoji and markdown.
    """
    if not text:
        return text
    text = _EMOJI_RE.sub("", text)
    text = _BOLD_MD_RE.sub(r"<b>\1</b>", text)
    text = _re.sub(r"  +", " ", text)
    return text.strip()


# ── Markdown table → reportlab Table conversion ─────────────────────────────
# The Decision Gates and Confidence Summary blocks are markdown tables. PDF's
# reportlab.Paragraph can't render those — it just shows the pipes literally.
# These helpers detect a markdown table inside a paragraph and convert it
# into a real reportlab Table so the PDF reads like a proper report.

def _is_markdown_table(text: str) -> bool:
    """A paragraph is a markdown table if it has ≥2 lines, all starting with |.
    Allows for an optional leading line (a caption, e.g. **Title**) above
    the table, so we check the longest run of |-lines."""
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    pipe_lines = [ln for ln in lines if ln.startswith("|")]
    return len(pipe_lines) >= 2


def _split_caption_and_table(text: str) -> tuple[str, str]:
    """Return (caption_text, table_block). Caption is anything before the
    first |-line; table_block is the contiguous run of |-lines."""
    lines = text.splitlines()
    caption_lines = []
    table_lines   = []
    seen_pipe     = False
    for ln in lines:
        s = ln.strip()
        if s.startswith("|"):
            seen_pipe = True
            table_lines.append(ln)
        elif seen_pipe:
            # Stop collecting table lines once we leave the pipe block
            break
        else:
            caption_lines.append(ln)
    return "\n".join(caption_lines).strip(), "\n".join(table_lines).strip()


def _parse_markdown_table(table_text: str) -> list[list[str]]:
    """Parse a markdown table into a 2D list of cell strings.
    The header-separator row (|---|---|) is dropped."""
    rows: list[list[str]] = []
    sep_pattern = _re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")
    for line in table_text.splitlines():
        if not line.strip():
            continue
        if sep_pattern.match(line):
            continue  # skip the |---|---| separator
        # Strip leading/trailing pipes, then split
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def _build_reportlab_table(rows, body_style, header_style, total_width):
    """Build a reportlab Table with sensible widths + alternating row shading."""
    from reportlab.platypus import Table, TableStyle, Paragraph
    from reportlab.lib import colors
    from reportlab.lib.units import inch

    if not rows:
        return None

    n_cols = max(len(r) for r in rows)

    # Pad short rows so the table doesn't crash on uneven row lengths
    rows = [r + [""] * (n_cols - len(r)) for r in rows]

    # Wrap every cell in a Paragraph so long text wraps correctly inside the
    # column. Header cells get the header style (bold + dark text on muted
    # background); body cells get the regular body style.
    data: list[list] = []
    for i, raw_row in enumerate(rows):
        row = []
        for cell in raw_row:
            cleaned = _pdf_clean(cell).replace("\n", "<br/>") or "—"
            style = header_style if i == 0 else body_style
            row.append(Paragraph(cleaned, style))
        data.append(row)

    # Column-width strategy: weight each column by the longest content it
    # contains, with a minimum floor so short-label columns don't get
    # crushed and long-prose columns get the room they need. This handles
    # both the 5-col Decision Gates table (1 prose column) and the 5-col
    # Confidence Summary table (3 prose columns) correctly.
    MIN_WEIGHT = 80   # ~1.2cm minimum per column
    col_max_lens = []
    for c in range(n_cols):
        longest = 0
        for r in range(len(rows)):
            if c < len(rows[r]):
                # Strip emoji + bold markers so they don't inflate weights
                clean = _pdf_clean(rows[r][c])
                longest = max(longest, len(clean))
        col_max_lens.append(max(MIN_WEIGHT, longest))
    total_weight = sum(col_max_lens) or 1
    col_widths = [(w / total_weight) * total_width for w in col_max_lens]

    tbl = Table(data, colWidths=col_widths, repeatRows=1, hAlign="LEFT")
    tbl.setStyle(TableStyle([
        ("BACKGROUND",   (0, 0), (-1, 0),  colors.HexColor("#1F2A44")),
        ("TEXTCOLOR",    (0, 0), (-1, 0),  colors.white),
        ("FONTNAME",     (0, 0), (-1, 0),  "Helvetica-Bold"),
        ("ALIGN",        (0, 0), (-1, -1), "LEFT"),
        ("VALIGN",       (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING",  (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING",   (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 6),
        ("GRID",         (0, 0), (-1, -1), 0.5, colors.HexColor("#D1D5DB")),
        # Alternating row shading on body rows (skip header)
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.white, colors.HexColor("#F9FAFB")]),
    ]))
    return tbl


def _render_pdf(title: str, sections: list[dict], out_path: Path) -> None:
    """Render a clean PDF using reportlab."""
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles    import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units     import inch
        from reportlab.lib           import colors
        from reportlab.platypus      import (
            SimpleDocTemplate, Paragraph, Spacer
        )
        from reportlab.lib.enums     import TA_CENTER
    except ImportError:
        print("  [report] reportlab not installed; writing plain-text .pdf fallback")
        out_path.write_text(_stitch_markdown(title, sections), encoding="utf-8")
        return

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "OAITitle", parent=styles["Title"],
        fontSize=22, leading=26,
        textColor=colors.HexColor("#1F2A44"),
        alignment=TA_CENTER, spaceAfter=4,
    )
    subtitle_style = ParagraphStyle(
        "OAISub", parent=styles["Italic"],
        fontSize=10, leading=12,
        textColor=colors.HexColor("#6B7280"),
        alignment=TA_CENTER, spaceAfter=24,
    )
    h2_style = ParagraphStyle(
        "OAIH2", parent=styles["Heading2"],
        fontSize=14, leading=18,
        textColor=colors.HexColor("#1F2A44"),
        spaceBefore=14, spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "OAIBody", parent=styles["BodyText"],
        fontSize=11, leading=15.5,
        textColor=colors.HexColor("#111827"),
        spaceAfter=8,
    )
    # Cell text style for body rows in tables — slightly tighter than the
    # main body so multi-row cells stay readable inside narrow columns.
    table_body_style = ParagraphStyle(
        "OAITableBody", parent=body_style,
        fontSize=9.5, leading=12.5, spaceAfter=0,
    )
    # Cell text style for the header row (white text on dark background)
    table_header_style = ParagraphStyle(
        "OAITableHeader", parent=body_style,
        fontSize=10, leading=12, spaceAfter=0,
        textColor=colors.white, fontName="Helvetica-Bold",
    )
    # Caption above a table (e.g. "**Decision Gate Summary**")
    table_caption_style = ParagraphStyle(
        "OAITableCaption", parent=body_style,
        fontSize=11, leading=14, spaceBefore=4, spaceAfter=4,
        textColor=colors.HexColor("#1F2A44"), fontName="Helvetica-Bold",
    )

    doc = SimpleDocTemplate(
        str(out_path), pagesize=LETTER,
        leftMargin=0.9*inch, rightMargin=0.9*inch,
        topMargin=0.9*inch,  bottomMargin=0.9*inch,
        title=title, author="OptimaAi",
    )

    # Width available for tables (page width minus side margins)
    table_total_width = LETTER[0] - 1.8 * inch

    story = []
    story.append(Paragraph(_pdf_clean(title), title_style))
    story.append(Paragraph(
        f"OptimaAi  &middot;  Generated {datetime.utcnow().strftime('%B %d, %Y')}",
        subtitle_style,
    ))

    for sec in sections:
        story.append(Paragraph(_pdf_clean(sec["title"]), h2_style))
        body = sec["content_md"].strip()
        for para in body.split("\n\n"):
            para = para.strip()
            if not para:
                continue

            # Markdown table detection — render as a real PDF Table
            if _is_markdown_table(para):
                caption, table_text = _split_caption_and_table(para)
                if caption:
                    story.append(Paragraph(
                        _pdf_clean(caption).replace("\n", "<br/>"),
                        table_caption_style,
                    ))
                rows = _parse_markdown_table(table_text)
                tbl  = _build_reportlab_table(
                    rows,
                    body_style   = table_body_style,
                    header_style = table_header_style,
                    total_width  = table_total_width,
                )
                if tbl is not None:
                    story.append(tbl)
                    story.append(Spacer(1, 0.10 * inch))
                continue

            # Regular prose — strip emoji, convert **bold** → <b>bold</b>
            cleaned = _pdf_clean(para).replace("\n", "<br/>")
            story.append(Paragraph(cleaned, body_style))
        story.append(Spacer(1, 0.15 * inch))

    doc.build(story)